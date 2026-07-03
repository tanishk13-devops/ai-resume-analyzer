import io
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

class TemplateRenderer:
    """Renders structured resume JSON into 25 beautifully styled, ATS-compliant formats (PDF, DOCX, TXT)"""

    # Dictionary of 25 custom template configurations
    TEMPLATES = {
        "executive_classic": {
            "name": "Executive Classic",
            "font": "Times-Roman",
            "font_bold": "Times-Bold",
            "primary": "#0f172a",
            "secondary": "#7f1d1d",
            "accent": "#475569",
            "align": "center",
            "divider_thickness": 1.0,
            "rgb_primary": RGBColor(15, 23, 42),
            "rgb_secondary": RGBColor(127, 29, 29)
        },
        "corporate_navy": {
            "name": "Corporate Navy",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#1e3a8a",
            "secondary": "#2563eb",
            "accent": "#64748b",
            "align": "left",
            "divider_thickness": 1.5,
            "rgb_primary": RGBColor(30, 58, 138),
            "rgb_secondary": RGBColor(37, 99, 235)
        },
        "tech_indigo": {
            "name": "Tech Indigo",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#312e81",
            "secondary": "#6d28d9",
            "accent": "#475569",
            "align": "left",
            "divider_thickness": 1.5,
            "rgb_primary": RGBColor(49, 46, 129),
            "rgb_secondary": RGBColor(109, 40, 217)
        },
        "emerald_minimalist": {
            "name": "Emerald Minimalist",
            "font": "Times-Roman",
            "font_bold": "Times-Bold",
            "primary": "#064e3b",
            "secondary": "#0d9488",
            "accent": "#64748b",
            "align": "center",
            "divider_thickness": 1.0,
            "rgb_primary": RGBColor(6, 78, 59),
            "rgb_secondary": RGBColor(13, 148, 136)
        },
        "modern_charcoal": {
            "name": "Modern Charcoal",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#334155",
            "secondary": "#64748b",
            "accent": "#94a3b8",
            "align": "left",
            "divider_thickness": 1.2,
            "rgb_primary": RGBColor(51, 65, 85),
            "rgb_secondary": RGBColor(100, 116, 139)
        },
        "executive_crimson": {
            "name": "Executive Crimson",
            "font": "Times-Roman",
            "font_bold": "Times-Bold",
            "primary": "#7f1d1d",
            "secondary": "#b45309",
            "accent": "#475569",
            "align": "center",
            "divider_thickness": 1.0,
            "rgb_primary": RGBColor(127, 29, 29),
            "rgb_secondary": RGBColor(180, 83, 9)
        },
        "clean_slate": {
            "name": "Clean Slate",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#0f172a",
            "secondary": "#64748b",
            "accent": "#94a3b8",
            "align": "left",
            "divider_thickness": 1.0,
            "rgb_primary": RGBColor(15, 23, 42),
            "rgb_secondary": RGBColor(100, 116, 139)
        },
        "silicon_valley": {
            "name": "Silicon Valley Tech",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#115e59",
            "secondary": "#0284c7",
            "accent": "#475569",
            "align": "left",
            "divider_thickness": 1.5,
            "rgb_primary": RGBColor(17, 94, 89),
            "rgb_secondary": RGBColor(2, 132, 199)
        },
        "academic_formal": {
            "name": "Academic Formal",
            "font": "Times-Roman",
            "font_bold": "Times-Bold",
            "primary": "#000000",
            "secondary": "#475569",
            "accent": "#64748b",
            "align": "center",
            "divider_thickness": 0.75,
            "rgb_primary": RGBColor(0, 0, 0),
            "rgb_secondary": RGBColor(71, 85, 105)
        },
        "finance_gold": {
            "name": "Finance Gold",
            "font": "Times-Roman",
            "font_bold": "Times-Bold",
            "primary": "#1e3a8a",
            "secondary": "#b45309",
            "accent": "#475569",
            "align": "center",
            "divider_thickness": 1.0,
            "rgb_primary": RGBColor(30, 58, 138),
            "rgb_secondary": RGBColor(180, 83, 9)
        },
        "startup_bold": {
            "name": "Startup Bold",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#4338ca",
            "secondary": "#db2777",
            "accent": "#475569",
            "align": "left",
            "divider_thickness": 2.0,
            "rgb_primary": RGBColor(67, 56, 202),
            "rgb_secondary": RGBColor(219, 39, 119)
        },
        "medical_teal": {
            "name": "Medical Teal",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#0f766e",
            "secondary": "#14b8a6",
            "accent": "#475569",
            "align": "left",
            "divider_thickness": 1.5,
            "rgb_primary": RGBColor(15, 118, 110),
            "rgb_secondary": RGBColor(20, 184, 166)
        },
        "sales_director": {
            "name": "Sales Director",
            "font": "Times-Roman",
            "font_bold": "Times-Bold",
            "primary": "#991b1b",
            "secondary": "#475569",
            "accent": "#64748b",
            "align": "center",
            "divider_thickness": 1.2,
            "rgb_primary": RGBColor(153, 27, 27),
            "rgb_secondary": RGBColor(71, 85, 105)
        },
        "creative_teal": {
            "name": "Creative Teal",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#0d9488",
            "secondary": "#16a34a",
            "accent": "#475569",
            "align": "left",
            "divider_thickness": 1.5,
            "rgb_primary": RGBColor(13, 148, 136),
            "rgb_secondary": RGBColor(22, 163, 74)
        },
        "midnight_slate": {
            "name": "Midnight Slate",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#1e293b",
            "secondary": "#334155",
            "accent": "#64748b",
            "align": "center",
            "divider_thickness": 1.0,
            "rgb_primary": RGBColor(30, 41, 59),
            "rgb_secondary": RGBColor(51, 65, 85)
        },
        "warm_amber": {
            "name": "Warm Amber",
            "font": "Times-Roman",
            "font_bold": "Times-Bold",
            "primary": "#451a03",
            "secondary": "#d97706",
            "accent": "#57534e",
            "align": "left",
            "divider_thickness": 1.2,
            "rgb_primary": RGBColor(69, 26, 3),
            "rgb_secondary": RGBColor(217, 119, 6)
        },
        "traditional_black": {
            "name": "Traditional Black",
            "font": "Times-Roman",
            "font_bold": "Times-Bold",
            "primary": "#000000",
            "secondary": "#000000",
            "accent": "#334155",
            "align": "center",
            "divider_thickness": 1.0,
            "rgb_primary": RGBColor(0, 0, 0),
            "rgb_secondary": RGBColor(0, 0, 0)
        },
        "bold_accent": {
            "name": "Bold Accent",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#312e81",
            "secondary": "#f97316",
            "accent": "#475569",
            "align": "left",
            "divider_thickness": 2.0,
            "rgb_primary": RGBColor(49, 46, 129),
            "rgb_secondary": RGBColor(249, 115, 22)
        },
        "simple_slate": {
            "name": "Simple Slate",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#475569",
            "secondary": "#64748b",
            "accent": "#94a3b8",
            "align": "left",
            "divider_thickness": 1.0,
            "rgb_primary": RGBColor(71, 85, 105),
            "rgb_secondary": RGBColor(100, 116, 139)
        },
        "metro_red": {
            "name": "Metro Red",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#be123c",
            "secondary": "#374151",
            "accent": "#6b7280",
            "align": "left",
            "divider_thickness": 1.5,
            "rgb_primary": RGBColor(190, 18, 60),
            "rgb_secondary": RGBColor(55, 65, 81)
        },
        "legal_eagle": {
            "name": "Legal Eagle",
            "font": "Times-Roman",
            "font_bold": "Times-Bold",
            "primary": "#1c1917",
            "secondary": "#44403c",
            "accent": "#78716c",
            "align": "center",
            "divider_thickness": 1.0,
            "rgb_primary": RGBColor(28, 25, 23),
            "rgb_secondary": RGBColor(68, 64, 60)
        },
        "product_lead": {
            "name": "Product Lead",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#0369a1",
            "secondary": "#0284c7",
            "accent": "#475569",
            "align": "left",
            "divider_thickness": 1.5,
            "rgb_primary": RGBColor(3, 105, 161),
            "rgb_secondary": RGBColor(2, 132, 199)
        },
        "data_scientist": {
            "name": "Data Scientist",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#14532d",
            "secondary": "#22c55e",
            "accent": "#475569",
            "align": "left",
            "divider_thickness": 1.5,
            "rgb_primary": RGBColor(20, 83, 45),
            "rgb_secondary": RGBColor(34, 197, 94)
        },
        "global_director": {
            "name": "Global Director",
            "font": "Times-Roman",
            "font_bold": "Times-Bold",
            "primary": "#1d4ed8",
            "secondary": "#d97706",
            "accent": "#475569",
            "align": "center",
            "divider_thickness": 1.2,
            "rgb_primary": RGBColor(29, 78, 216),
            "rgb_secondary": RGBColor(217, 119, 6)
        },
        "hr_generalist": {
            "name": "HR Generalist",
            "font": "Helvetica",
            "font_bold": "Helvetica-Bold",
            "primary": "#9f1239",
            "secondary": "#e11d48",
            "accent": "#64748b",
            "align": "left",
            "divider_thickness": 1.2,
            "rgb_primary": RGBColor(159, 18, 57),
            "rgb_secondary": RGBColor(225, 29, 72)
        }
    }

    def get_template_keys(self):
        """Return list of all template keys and their display names"""
        return [(k, v["name"]) for k, v in self.TEMPLATES.items()]

    def generate_pdf(self, data, template_id="executive_classic", fit_single_page=False, font_size_modifier=0.0, spacing_modifier=0.0):
        """Generate a beautifully formatted PDF based on a template configuration, with an optional fit_single_page layout compression."""
        buffer = io.BytesIO()
        config = self.TEMPLATES.get(template_id, self.TEMPLATES["executive_classic"])
        
        # Colors conversion
        primary_color = colors.HexColor(config["primary"])
        secondary_color = colors.HexColor(config["secondary"])
        accent_color = colors.HexColor(config["accent"])
        
        # Base modifiers: if fit_single_page is True, apply default compression if modifiers are 0
        f_mod = font_size_modifier
        s_mod = spacing_modifier
        if fit_single_page and f_mod == 0.0 and s_mod == 0.0:
            f_mod = -1.0
            s_mod = -2.0
            
        # Margins: 0.75" (54 pt) or compact 0.4" (30 pt)
        margin_size = max(20, 30 if fit_single_page else (54 + s_mod * 4))
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter,
            rightMargin=margin_size, 
            leftMargin=margin_size,
            topMargin=margin_size, 
            bottomMargin=margin_size
        )
        
        story = []
        align_code = 1 if config["align"] == "center" else 0

        # Styles definition (compressed font size/leading if fit_single_page or modifiers are set)
        name_size = max(12, 16 if fit_single_page else (22 + f_mod * 1.5))
        name_leading = name_size + 4
        name_space = max(2, 4 if fit_single_page else (6 + s_mod))
        
        name_style = ParagraphStyle(
            'ResName',
            fontName=config["font_bold"],
            fontSize=name_size,
            leading=name_leading,
            textColor=primary_color,
            alignment=align_code,
            spaceAfter=name_space
        )
        
        contact_size = max(7.0, 8.5 if fit_single_page else (9.5 + f_mod))
        contact_leading = contact_size + 2
        contact_space = max(2, 6 if fit_single_page else (12 + s_mod * 1.5))
        
        contact_style = ParagraphStyle(
            'ResContact',
            fontName=config["font"],
            fontSize=contact_size,
            leading=contact_leading,
            textColor=accent_color,
            alignment=align_code,
            spaceAfter=contact_space
        )
        
        h1_size = max(8.5, 9.5 if fit_single_page else (11 + f_mod))
        h1_leading = h1_size + 2
        h1_space_before = max(2, 6 if fit_single_page else (12 + s_mod * 1.5))
        h1_space_after = max(1, 2 if fit_single_page else (4 + s_mod * 0.5))
        
        h1_style = ParagraphStyle(
            'ResH1',
            fontName=config["font_bold"],
            fontSize=h1_size,
            leading=h1_leading,
            textColor=primary_color,
            spaceBefore=h1_space_before,
            spaceAfter=h1_space_after,
            keepWithNext=True
        )
        
        body_size = max(7.5, 8.5 if fit_single_page else (9.5 + f_mod))
        body_leading = body_size + 3.5
        body_space_after = max(1, 3 if fit_single_page else (5 + s_mod * 0.5))
        
        body_style = ParagraphStyle(
            'ResBody',
            fontName=config["font"],
            fontSize=body_size,
            leading=body_leading,
            textColor=colors.HexColor("#1e293b"),
            spaceAfter=body_space_after
        )
        
        bullet_size = max(7.0, 8.0 if fit_single_page else (9.0 + f_mod))
        bullet_leading = bullet_size + 3.0
        bullet_space_after = max(1, 1.5 if fit_single_page else (3 + s_mod * 0.5))
        
        bullet_style = ParagraphStyle(
            'ResBullet',
            fontName=config["font"],
            fontSize=bullet_size,
            leading=bullet_leading,
            textColor=colors.HexColor("#1e293b"),
            leftIndent=15,
            firstLineIndent=-10,
            spaceAfter=bullet_space_after
        )
        
        # 1. Header (Name & Contact)
        story.append(Paragraph(data.get("name", "Resume Owner"), name_style))
        
        contact_data = data.get("contact", {})
        contact_parts = []
        if contact_data.get("email"): contact_parts.append(contact_data["email"])
        if contact_data.get("phone"): contact_parts.append(contact_data["phone"])
        if contact_data.get("linkedin"): contact_parts.append(contact_data["linkedin"])
        if contact_data.get("location"): contact_parts.append(contact_data["location"])
            
        contact_str = "  |  ".join(contact_parts)
        story.append(Paragraph(contact_str, contact_style))
        
        # Divider Line
        div_space_after = max(2, 4 if fit_single_page else (8 + s_mod))
        story.append(HRFlowable(width="100%", thickness=config["divider_thickness"], color=secondary_color, spaceAfter=div_space_after))
        
        # 2. Summary
        summary = data.get("summary")
        if summary:
            story.append(Paragraph("SUMMARY", h1_style))
            hdr_space_after = max(2, 4 if fit_single_page else (6 + s_mod * 0.5))
            story.append(HRFlowable(width="100%", thickness=0.5, color=accent_color, spaceAfter=hdr_space_after))
            story.append(Paragraph(summary, body_style))
            story.append(Spacer(1, max(1, 2 if fit_single_page else (4 + s_mod * 0.5))))
            
        # 3. Work Experience
        experience = data.get("experience", [])
        if experience:
            story.append(Paragraph("PROFESSIONAL EXPERIENCE", h1_style))
            hdr_space_after = max(2, 4 if fit_single_page else (8 + s_mod))
            story.append(HRFlowable(width="100%", thickness=0.5, color=accent_color, spaceAfter=hdr_space_after))
            
            # If margins are small, printable width increases
            table_col_widths = [390, 160] if (fit_single_page or s_mod < 0) else [350, 150]
            
            for exp in experience:
                job_title = exp.get("job_title", "")
                company = exp.get("company", "")
                dates = exp.get("dates", "")
                
                header_data = [
                    [
                        Paragraph(f"<b>{job_title}</b>", body_style), 
                        Paragraph(f"<font color='{accent_color.hexval()}'><b>{dates}</b></font>", ParagraphStyle('RAlign', parent=body_style, alignment=2))
                    ],
                    [
                        Paragraph(f"<i>{company}</i>", body_style), 
                        Paragraph("", body_style)
                    ]
                ]
                
                header_table = Table(header_data, colWidths=table_col_widths)
                header_table.setStyle(TableStyle([
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('LEFTPADDING', (0,0), (-1,-1), 0),
                    ('RIGHTPADDING', (0,0), (-1,-1), 0),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                    ('TOPPADDING', (0,0), (-1,-1), 0),
                ]))
                story.append(header_table)
                story.append(Spacer(1, max(0.5, 1 if fit_single_page else (2 + s_mod * 0.25))))
                
                # Bullets
                for bullet in exp.get("description", []):
                    if bullet.strip():
                        story.append(Paragraph(f"&bull; {bullet.strip()}", bullet_style))
                story.append(Spacer(1, max(1, 2 if fit_single_page else (5 + s_mod * 0.5))))

        # 4. Skills
        skills = data.get("skills", [])
        if skills:
            story.append(Paragraph("TECHNICAL SKILLS", h1_style))
            hdr_space_after = max(2, 4 if fit_single_page else (6 + s_mod * 0.5))
            story.append(HRFlowable(width="100%", thickness=0.5, color=accent_color, spaceAfter=hdr_space_after))
            skills_str = ", ".join(skills)
            story.append(Paragraph(skills_str, body_style))
            story.append(Spacer(1, max(1, 2 if fit_single_page else (4 + s_mod * 0.5))))
            
        # 5. Education
        education = data.get("education", [])
        if education:
            story.append(Paragraph("EDUCATION", h1_style))
            hdr_space_after = max(2, 4 if fit_single_page else (8 + s_mod))
            story.append(HRFlowable(width="100%", thickness=0.5, color=accent_color, spaceAfter=hdr_space_after))
            
            table_col_widths_edu = [420, 130] if (fit_single_page or s_mod < 0) else [380, 120]
            
            for edu in education:
                degree = edu.get("degree", "")
                school = edu.get("school", "")
                dates = edu.get("dates", "")
                
                edu_data = [
                    [
                        Paragraph(f"<b>{degree}</b> - <i>{school}</i>", body_style),
                        Paragraph(f"<font color='{accent_color.hexval()}'><b>{dates}</b></font>", ParagraphStyle('RAlign', parent=body_style, alignment=2))
                    ]
                ]
                edu_table = Table(edu_data, colWidths=table_col_widths_edu)
                edu_table.setStyle(TableStyle([
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('LEFTPADDING', (0,0), (-1,-1), 0),
                    ('RIGHTPADDING', (0,0), (-1,-1), 0),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                    ('TOPPADDING', (0,0), (-1,-1), 0),
                ]))
                story.append(edu_table)
                story.append(Spacer(1, max(1, 2 if fit_single_page else (4 + s_mod * 0.5))))
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_docx(self, data, template_id="executive_classic", fit_single_page=False, font_size_modifier=0.0, spacing_modifier=0.0):
        """Generate a styled DOCX based on a template configuration, with optional fit_single_page layout compression."""
        doc = docx.Document()
        config = self.TEMPLATES.get(template_id, self.TEMPLATES["executive_classic"])
        
        # Base modifiers: if fit_single_page is True, apply default compression if modifiers are 0
        f_mod = font_size_modifier
        s_mod = spacing_modifier
        if fit_single_page and f_mod == 0.0 and s_mod == 0.0:
            f_mod = -1.0
            s_mod = -2.0
            
        # Margins: 0.75" or compact 0.4"
        margin_size = max(0.25, 0.4 if fit_single_page else (0.75 + s_mod * 0.05))
        for section in doc.sections:
            section.top_margin = Inches(margin_size)
            section.bottom_margin = Inches(margin_size)
            section.left_margin = Inches(margin_size)
            section.right_margin = Inches(margin_size)
            
        font_name = "Times New Roman" if config["font"] == "Times-Roman" else "Arial"
        align_code = WD_ALIGN_PARAGRAPH.CENTER if config["align"] == "center" else WD_ALIGN_PARAGRAPH.LEFT
        
        # 1. Header Name
        p_name = doc.add_paragraph()
        p_name.alignment = align_code
        run_name = p_name.add_run(data.get("name", "Resume Owner"))
        run_name.font.name = font_name
        run_name.font.size = Pt(max(12, 16 if fit_single_page else (22 + f_mod * 1.5)))
        run_name.bold = True
        run_name.font.color.rgb = config["rgb_primary"]
        
        # Contact Details
        contact = data.get("contact", {})
        contact_parts = []
        if contact.get("email"): contact_parts.append(contact["email"])
        if contact.get("phone"): contact_parts.append(contact["phone"])
        if contact.get("linkedin"): contact_parts.append(contact["linkedin"])
        if contact.get("location"): contact_parts.append(contact["location"])
        
        p_contact = doc.add_paragraph()
        p_contact.alignment = align_code
        run_contact = p_contact.add_run("  |  ".join(contact_parts))
        run_contact.font.name = font_name
        run_contact.font.size = Pt(max(7.0, 8.5 if fit_single_page else (9.5 + f_mod)))
        run_contact.font.color.rgb = config["rgb_secondary"]
        
        # Divider Line
        def add_section_heading(title):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(max(2, 6 if fit_single_page else (12 + s_mod * 1.5)))
            p.paragraph_format.space_after = Pt(max(0.5, 1 if fit_single_page else (2 + s_mod * 0.25)))
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(title.upper())
            run.font.name = font_name
            run.font.size = Pt(max(8.5, 9.5 if fit_single_page else (11 + f_mod)))
            run.bold = True
            run.font.color.rgb = config["rgb_primary"]
            
            # Bottom border line
            p_border = doc.add_paragraph()
            p_border.paragraph_format.space_after = Pt(max(1, 3 if fit_single_page else (6 + s_mod * 0.5)))
            # Adjust divider length to match margins
            border_len = 95 if fit_single_page or s_mod < 0 else 70
            run_border = p_border.add_run("―" * border_len)
            run_border.font.size = Pt(4)
            run_border.font.color.rgb = config["rgb_secondary"]
            
        # 2. Summary
        summary = data.get("summary")
        if summary:
            add_section_heading("Summary")
            p_sum = doc.add_paragraph()
            p_sum.paragraph_format.space_after = Pt(max(1, 3 if fit_single_page else (6 + s_mod * 0.5)))
            run_sum = p_sum.add_run(summary)
            run_sum.font.name = font_name
            run_sum.font.size = Pt(max(7.5, 8.5 if fit_single_page else (9.5 + f_mod)))
            
        # 3. Work Experience
        experience = data.get("experience", [])
        if experience:
            add_section_heading("Professional Experience")
            for exp in experience:
                p_header = doc.add_paragraph()
                p_header.paragraph_format.space_before = Pt(max(1, 2 if fit_single_page else (4 + s_mod * 0.5)))
                p_header.paragraph_format.space_after = Pt(0.5)
                
                # Align dates to the right using native tab stops
                from docx.enum.text import WD_TAB_ALIGNMENT
                # Printable page width is 8.5" minus margins
                right_margin_inch = 8.5 - 2 * margin_size
                p_header.paragraph_format.tab_stops.add_tab_stop(Inches(right_margin_inch), WD_TAB_ALIGNMENT.RIGHT)
                
                run_title = p_header.add_run(f"{exp.get('job_title', '')}")
                run_title.font.name = font_name
                run_title.font.size = Pt(max(7.5, 8.5 if fit_single_page else (9.5 + f_mod)))
                run_title.bold = True
                
                p_header.add_run("\t")
                
                run_dates = p_header.add_run(exp.get('dates', ''))
                run_dates.font.name = font_name
                run_dates.font.size = Pt(max(7.5, 8.5 if fit_single_page else (9.5 + f_mod)))
                run_dates.bold = True
                run_dates.font.color.rgb = config["rgb_secondary"]
                
                p_sub = doc.add_paragraph()
                p_sub.paragraph_format.space_after = Pt(max(1, 2 if fit_single_page else (4 + s_mod * 0.5)))
                run_comp = p_sub.add_run(exp.get('company', ''))
                run_comp.font.name = font_name
                run_comp.font.size = Pt(max(7.5, 8.5 if fit_single_page else (9.5 + f_mod)))
                run_comp.italic = True
                
                # Bullets
                for bullet in exp.get("description", []):
                    if bullet.strip():
                        p_bullet = doc.add_paragraph(style='List Bullet')
                        p_bullet.paragraph_format.space_after = Pt(max(0.5, 1.5 if fit_single_page else (2.5 + s_mod * 0.25)))
                        p_bullet.paragraph_format.left_indent = Inches(0.20 if fit_single_page or s_mod < 0 else 0.25)
                        run_b = p_bullet.add_run(bullet.strip())
                        run_b.font.name = font_name
                        run_b.font.size = Pt(max(7.0, 8.5 if fit_single_page else (9.5 + f_mod)))
                        
        # 4. Skills
        skills = data.get("skills", [])
        if skills:
            add_section_heading("Technical Skills")
            p_skills = doc.add_paragraph()
            p_skills.paragraph_format.space_after = Pt(max(1, 3 if fit_single_page else (6 + s_mod * 0.5)))
            run_skills = p_skills.add_run(", ".join(skills))
            run_skills.font.name = font_name
            run_skills.font.size = Pt(max(7.5, 8.5 if fit_single_page else (9.5 + f_mod)))
            
        # 5. Education
        education = data.get("education", [])
        if education:
            add_section_heading("Education")
            for edu in education:
                p_edu = doc.add_paragraph()
                p_edu.paragraph_format.space_after = Pt(max(1, 2 if fit_single_page else (3 + s_mod * 0.25)))
                
                # Align dates to the right using native tab stops
                from docx.enum.text import WD_TAB_ALIGNMENT
                right_margin_inch = 8.5 - 2 * margin_size
                p_edu.paragraph_format.tab_stops.add_tab_stop(Inches(right_margin_inch), WD_TAB_ALIGNMENT.RIGHT)
                
                run_deg = p_edu.add_run(f"{edu.get('degree', '')} - {edu.get('school', '')}")
                run_deg.font.name = font_name
                run_deg.font.size = Pt(max(7.5, 8.5 if fit_single_page else (9.5 + f_mod)))
                run_deg.bold = True
                
                p_edu.add_run("\t")
                
                run_dates = p_edu.add_run(edu.get('dates', ''))
                run_dates.font.name = font_name
                run_dates.font.size = Pt(max(7.5, 8.5 if fit_single_page else (9.5 + f_mod)))
                run_dates.bold = True
                run_dates.font.color.rgb = config["rgb_secondary"]

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()

    def generate_txt(self, data):
        """Generate a standard formatted text layout"""
        out = []
        out.append(data.get("name", "Resume Owner").upper())
        
        contact = data.get("contact", {})
        contact_parts = []
        if contact.get("email"): contact_parts.append(contact["email"])
        if contact.get("phone"): contact_parts.append(contact["phone"])
        if contact.get("linkedin"): contact_parts.append(contact["linkedin"])
        if contact.get("location"): contact_parts.append(contact["location"])
        out.append(" | ".join(contact_parts))
        out.append("=" * 60)
        out.append("")
        
        summary = data.get("summary")
        if summary:
            out.append("SUMMARY")
            out.append("-" * 30)
            out.append(summary)
            out.append("")
            
        experience = data.get("experience", [])
        if experience:
            out.append("PROFESSIONAL EXPERIENCE")
            out.append("-" * 30)
            for exp in experience:
                out.append(f"{exp.get('job_title', '')} | {exp.get('company', '')}")
                out.append(exp.get('dates', ''))
                for bullet in exp.get("description", []):
                    if bullet.strip():
                        out.append(f"  - {bullet.strip()}")
                out.append("")
                
        skills = data.get("skills", [])
        if skills:
            out.append("TECHNICAL SKILLS")
            out.append("-" * 30)
            out.append(", ".join(skills))
            out.append("")
            
        education = data.get("education", [])
        if education:
            out.append("EDUCATION")
            out.append("-" * 30)
            for edu in education:
                out.append(f"{edu.get('degree', '')} | {edu.get('school', '')} ({edu.get('dates', '')})")
                out.append("")
                
        return "\n".join(out)

    def generate_cover_letter_pdf(self, letter_text, metadata, template_id="executive_classic"):
        """Generate a styled Cover Letter PDF using the selected resume template's branding style"""
        buffer = io.BytesIO()
        config = self.TEMPLATES.get(template_id, self.TEMPLATES["executive_classic"])
        
        # Colors conversion
        primary_color = colors.HexColor(config["primary"])
        secondary_color = colors.HexColor(config["secondary"])
        accent_color = colors.HexColor(config["accent"])
        
        # 0.75" margins
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter,
            rightMargin=54, 
            leftMargin=54,
            topMargin=54, 
            bottomMargin=54
        )
        
        story = []
        align_code = 1 if config["align"] == "center" else 0

        # Styles definition
        name_style = ParagraphStyle(
            'LetterName',
            fontName=config["font_bold"],
            fontSize=22,
            leading=26,
            textColor=primary_color,
            alignment=align_code,
            spaceAfter=6
        )
        
        contact_style = ParagraphStyle(
            'LetterContact',
            fontName=config["font"],
            fontSize=9.5,
            leading=12,
            textColor=accent_color,
            alignment=align_code,
            spaceAfter=12
        )
        
        body_style = ParagraphStyle(
            'LetterBody',
            fontName=config["font"],
            fontSize=10,
            leading=14.5,
            textColor=colors.HexColor("#1e293b"),
            spaceAfter=12
        )
        
        # 1. Header (Name & Contact info matching resume)
        story.append(Paragraph(metadata.get("name", "Applicant Name"), name_style))
        
        contact_parts = []
        if metadata.get("email"): contact_parts.append(metadata["email"])
        if metadata.get("phone"): contact_parts.append(metadata["phone"])
        if metadata.get("linkedin"): contact_parts.append(metadata["linkedin"])
        if metadata.get("location"): contact_parts.append(metadata["location"])
            
        contact_str = "  |  ".join(contact_parts)
        story.append(Paragraph(contact_str, contact_style))
        
        # Divider Line
        story.append(HRFlowable(width="100%", thickness=config["divider_thickness"], color=secondary_color, spaceAfter=18))

        # 2. Date & Address metadata (if any)
        date_str = metadata.get("date", "")
        if not date_str:
            from datetime import datetime
            date_str = datetime.now().strftime("%B %d, %Y")
            
        story.append(Paragraph(f"<b>{date_str}</b>", body_style))
        
        recipient_parts = []
        if metadata.get("recipient_name"): recipient_parts.append(metadata["recipient_name"])
        if metadata.get("recipient_title"): recipient_parts.append(metadata["recipient_title"])
        if metadata.get("company_name"): recipient_parts.append(metadata["company_name"])
        
        if recipient_parts:
            rec_text = "<br/>".join(recipient_parts)
            story.append(Paragraph(rec_text, body_style))
            story.append(Spacer(1, 6))

        # 3. Letter Body (split by newlines to keep structure)
        # Clean double newlines and parse paragraphs
        paragraphs = [p.strip() for p in letter_text.split('\n') if p.strip()]
        
        for p in paragraphs:
            # Check if it starts with Salutation or Sign-off
            story.append(Paragraph(p, body_style))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_cover_letter_docx(self, letter_text, metadata, template_id="executive_classic"):
        """Generate a styled Cover Letter DOCX using the selected resume template's branding style"""
        doc = docx.Document()
        config = self.TEMPLATES.get(template_id, self.TEMPLATES["executive_classic"])
        
        # Margins: 0.75" everywhere
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
        font_name = "Times New Roman" if config["font"] == "Times-Roman" else "Arial"
        align_code = WD_ALIGN_PARAGRAPH.CENTER if config["align"] == "center" else WD_ALIGN_PARAGRAPH.LEFT
        
        # 1. Header Name
        p_name = doc.add_paragraph()
        p_name.alignment = align_code
        run_name = p_name.add_run(metadata.get("name", "Applicant Name"))
        run_name.font.name = font_name
        run_name.font.size = Pt(22)
        run_name.bold = True
        run_name.font.color.rgb = config["rgb_primary"]
        
        # Contact Details
        contact_parts = []
        if metadata.get("email"): contact_parts.append(metadata["email"])
        if metadata.get("phone"): contact_parts.append(metadata["phone"])
        if metadata.get("linkedin"): contact_parts.append(metadata["linkedin"])
        if metadata.get("location"): contact_parts.append(metadata["location"])
        
        p_contact = doc.add_paragraph()
        p_contact.alignment = align_code
        run_contact = p_contact.add_run("  |  ".join(contact_parts))
        run_contact.font.name = font_name
        run_contact.font.size = Pt(9.5)
        run_contact.font.color.rgb = config["rgb_secondary"]
        
        # Divider Line (series of characters matching resume)
        p_border = doc.add_paragraph()
        p_border.paragraph_format.space_before = Pt(4)
        p_border.paragraph_format.space_after = Pt(12)
        run_border = p_border.add_run("―" * 70)
        run_border.font.size = Pt(4)
        run_border.font.color.rgb = config["rgb_secondary"]

        # Date & Metadata
        date_str = metadata.get("date", "")
        if not date_str:
            from datetime import datetime
            date_str = datetime.now().strftime("%B %d, %Y")
            
        p_date = doc.add_paragraph()
        p_date.paragraph_format.space_after = Pt(6)
        run_date = p_date.add_run(date_str)
        run_date.font.name = font_name
        run_date.font.size = Pt(10)
        run_date.bold = True
        
        recipient_parts = []
        if metadata.get("recipient_name"): recipient_parts.append(metadata["recipient_name"])
        if metadata.get("recipient_title"): recipient_parts.append(metadata["recipient_title"])
        if metadata.get("company_name"): recipient_parts.append(metadata["company_name"])
        
        if recipient_parts:
            p_rec = doc.add_paragraph()
            p_rec.paragraph_format.space_after = Pt(12)
            run_rec = p_rec.add_run("\n".join(recipient_parts))
            run_rec.font.name = font_name
            run_rec.font.size = Pt(10)

        # Body paragraphs
        paragraphs = [p.strip() for p in letter_text.split('\n') if p.strip()]
        for p in paragraphs:
            p_body = doc.add_paragraph()
            p_body.paragraph_format.space_after = Pt(10)
            run_body = p_body.add_run(p)
            run_body.font.name = font_name
            run_body.font.size = Pt(10)
            
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()
