import os
import json
from pypdf import PdfReader
from docx import Document
from openai import OpenAI

def extract_text_from_file(uploaded_file) -> str:
    """
    Extracts text content from a Streamlit UploadedFile object.
    Supports PDF, DOCX, and TXT files.
    """
    filename = uploaded_file.name.lower()
    text = ""
    
    try:
        if filename.endswith(".pdf"):
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        elif filename.endswith(".docx"):
            doc = Document(uploaded_file)
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        elif filename.endswith(".txt"):
            text = uploaded_file.read().decode("utf-8", errors="ignore")
        else:
            raise ValueError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")
    except Exception as e:
        raise RuntimeError(f"Error reading file {filename}: {str(e)}")
        
    return text.strip()

def get_mock_analysis(resume_text: str, job_description: str) -> dict:
    """
    Returns a highly realistic backend engineering mock analysis to demonstrate all features.
    """
    return {
        "match_score": 68,
        "summary": "The candidate shows strong foundational skills in Python, JavaScript, and general software development, making them a moderate fit for the position. However, to align with the 'Senior/Mid-level Backend Engineer' requirements, the resume needs to emphasize experience with backend web frameworks (such as FastAPI or Django), database optimizations, containerization (Docker), and cloud infrastructure setup (AWS). The current bullet points are also mostly task-based rather than achievement-oriented.",
        "strengths": [
            "Good core foundation in Python and SQL database structures.",
            "Versatile development experience spanning both frontend (HTML, CSS, JS) and backend scripts.",
            "Relevant academic background with a B.S. in Computer Science."
        ],
        "keyword_gap_analysis": {
            "matching_keywords": ["Python", "JavaScript", "HTML", "CSS", "SQL", "Git", "AWS"],
            "missing_keywords": ["FastAPI", "PostgreSQL", "Docker", "CI/CD", "Redis", "Kafka", "REST APIs", "GraphQL"]
        },
        "bullet_point_improvements": [
            {
                "original": "Worked on developing the backend API.",
                "improved": "Designed and implemented 10+ scalable REST API endpoints using Python, reducing client latency by 15%.",
                "rationale": "Introduced active verbs, quantified results, and highlighted API architecture alignment with backend expectations."
            },
            {
                "original": "Wrote databases queries to retrieve data.",
                "improved": "Optimized complex SQL queries and structured indexing patterns, improving query response speeds by 30% on PostgreSQL databases.",
                "rationale": "Specified database system (PostgreSQL) and focused on performance engineering which is highly valued in backend roles."
            },
            {
                "original": "Fixed bugs and helped deploy releases.",
                "improved": "Streamlined git-based releases and automated deployments, decreasing deployment-related downtime by 12%.",
                "rationale": "Highlights familiarity with modern devops workflow and proactive deployment responsibility instead of passive 'helping'."
            }
        ],
        "tailored_sections": [
            {
                "section_name": "Professional Summary",
                "content": "Backend Engineer with 3+ years of experience designing API architectures, optimizing database queries, and managing cloud deployments. Proficient in Python, SQL, and Git workflows, with a proven track record of reducing latency and improving server-side performance. Seeking to leverage backend skills to build high-performance APIs at your scale."
            },
            {
                "section_name": "Technical Skills (Optimized)",
                "content": "Languages: Python, SQL (PostgreSQL), JavaScript, HTML/CSS\nFrameworks & Libraries: FastAPI (Familiar), REST APIs, GraphQL\nTools & DevOps: Git, Docker, AWS (EC2, RDS, Route53), CI/CD"
            }
        ],
        "interview_questions": [
            {
                "question": "The job description mentions FastAPI and PostgreSQL. Can you walk us through how you would optimize a slow-performing database query in a FastAPI app?",
                "answer_strategy": "Explain your step-by-step optimization process: using EXPLAIN ANALYZE to identify sequential scans, adding proper indexes (B-tree, GIN), refactoring ORM queries to avoid N+1 issues, and implementing connection pooling in FastAPI using AsyncSession.",
                "sample_answer": "I would start by profiling the query with PostgreSQL's 'EXPLAIN ANALYZE' to check the execution plan. If I see sequential scans on large tables, I would introduce targeted indexing. In the FastAPI application, I would verify that we are using SQLAlchemy's async connection pool to avoid blocking I/O, and refactor queries using options like 'joinedload' to eliminate N+1 select issues."
            },
            {
                "question": "Describe your experience with containerization and managing Docker deployments.",
                "answer_strategy": "Discuss writing a clean Dockerfile (multi-stage builds to minimize image size), managing multi-container setups via Docker Compose, and deploying them to cloud platforms (like AWS ECS or EC2).",
                "sample_answer": "In my previous projects, I containerized the backend services using multi-stage Dockerfiles to build lightweight images. I configured Docker Compose for local multi-service testing (API, database, Redis cache) and pushed the tested images to AWS ECR, which were then deployed on AWS EC2 instances, ensuring identical environments between dev and prod."
            }
        ]
    }

def analyze_resume(resume_text: str, job_description: str, api_key: str, provider: str = "OpenAI") -> dict:
    """
    Performs a detailed, structured analysis comparing the resume and job description.
    Supports both OpenAI and Google Gemini (via its OpenAI-compatible endpoint).
    """
    # Check if simulation/demo key is supplied
    if api_key == "DEMO" or provider == "Demo Mode":
        return get_mock_analysis(resume_text, job_description)
        
    if not api_key:
        raise ValueError(f"{provider} API key is required. Please enter it in the sidebar.")
        
    # Configure client based on the selected AI provider
    if provider == "Google Gemini":
        client = OpenAI(
            api_key=api_key,
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
        )
        model_name = "gemini-2.5-flash"  # Highly compatible, fast, and supports JSON output
    else:
        client = OpenAI(api_key=api_key)
        model_name = "gpt-4o"
    
    system_prompt = (
        "You are an expert ATS (Applicant Tracking System) specialist and executive career coach. "
        "Your task is to analyze the provided resume text against the job description text and "
        "provide structured, actionable feedback to improve the resume for this specific job.\n\n"
        "You must respond ONLY with a single JSON object. Do not include any explanation before or "
        "after the JSON block. The JSON object must strictly match the following schema:\n"
        "{\n"
        '  "match_score": integer (0 to 100 representing how well the resume matches the job description),\n'
        '  "summary": "string (a professional, constructive summary of the candidate\'s fit for the role)",\n'
        '  "strengths": ["string (bullet point highlighting a key strength/relevant experience)"],\n'
        '  "keyword_gap_analysis": {\n'
        '    "matching_keywords": ["string (keywords or skills present in both resume and job description)"],\n'
        '    "missing_keywords": ["string (important keywords or skills in the job description that are missing or weak in the resume)"]\n'
        '  },\n'
        '  "bullet_point_improvements": [\n'
        '    {\n'
        '      "original": "string (the original weak or unoptimized bullet point from the resume)",\n'
        '      "improved": "string (the rewritten, highly impactful, action-oriented, and tailored bullet point using the CAR/STAR method and keywords from the job description)",\n'
        '      "rationale": "string (explanation of why the improvement is better and how it aligns with the JD)"\n'
        '    }\n'
        '  ],\n'
        '  "tailored_sections": [\n'
        '    {\n'
        '      "section_name": "string (name of the section, e.g., Professional Summary, Core Competencies, Projects)",\n'
        '      "content": "string (rewritten section content optimized for the target job description)"\n'
        '    }\n'
        '  ],\n'
        '  "interview_questions": [\n'
        '    {\n'
        '      "question": "string (a standard or technical interview question likely to be asked for this role based on the candidate\'s gaps or listed experience)",\n'
        '      "answer_strategy": "string (brief advice on how the candidate should approach answering this question)",\n'
        '      "sample_answer": "string (a high-quality sample answer tailored to the candidate\'s profile)"\n'
        '    }\n'
        '  ]\n'
        "}\n\n"
        "Ensure all keys and structures are present. Be highly constructive, critical, and objective."
    )
    
    user_prompt = (
        f"--- RESUME TEXT ---\n{resume_text}\n\n"
        f"--- JOB DESCRIPTION ---\n{job_description}\n\n"
        f"Please analyze the resume against the job description and output the JSON response."
    )
    
    try:
        response = client.chat.completions.create(
            model=model_name,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2, # low temperature for consistency and objectivity
        )
        
        result_content = response.choices[0].message.content
        return json.loads(result_content)
    except json.JSONDecodeError as je:
        raise RuntimeError(f"Failed to parse AI response as valid JSON: {str(je)}")
    except Exception as e:
        raise RuntimeError(f"{provider} API Error: {str(e)}")
