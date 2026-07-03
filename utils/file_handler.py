import PyPDF2
import docx
import io

class FileHandler:
    """Handle file uploads and text extraction"""
    
    ALLOWED_EXTENSIONS = {'.pdf', '.docx'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx']
    
    def extract_text(self, uploaded_file):
        """
        Extract text from uploaded file
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            str: Extracted text from file
        """
        if uploaded_file is None:
            raise ValueError("No file provided")
        
        # Check file size
        if uploaded_file.size > self.MAX_FILE_SIZE:
            raise ValueError(f"File size exceeds maximum of 10MB")
        
        # Get file extension
        file_ext = uploaded_file.name.split('.')[-1].lower()
        
        if file_ext == 'pdf':
            return self._extract_from_pdf(uploaded_file)
        elif file_ext == 'docx':
            return self._extract_from_docx(uploaded_file)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_from_pdf(self, pdf_file):
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    def _extract_from_docx(self, docx_file):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(docx_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
    
    def validate_file(self, uploaded_file):
        """Validate uploaded file"""
        if uploaded_file is None:
            return False, "No file provided"
        
        if uploaded_file.size == 0:
            return False, "File is empty"
        
        if uploaded_file.size > self.MAX_FILE_SIZE:
            return False, "File size exceeds 10MB limit"
        
        file_ext = uploaded_file.name.split('.')[-1].lower()
        if file_ext not in self.supported_formats:
            return False, f"Unsupported format: {file_ext}"
        
        return True, "File is valid"

    def replace_text_in_docx(self, doc, target_text, replacement_text):
        """
        Replace target_text with replacement_text in a python-docx Document object,
        preserving original formatting and styles as much as possible.
        """
        if not doc or not target_text or not replacement_text:
            return False
            
        modified = False
        
        # Helper to clean target text for comparison
        clean_target = " ".join(target_text.strip().split())
        
        # 1. Search in paragraphs
        for paragraph in doc.paragraphs:
            clean_para_text = " ".join(paragraph.text.strip().split())
            if clean_target in clean_para_text or target_text in paragraph.text:
                # If target_text matches a run exactly, we replace in run
                for run in paragraph.runs:
                    if target_text in run.text:
                        run.text = run.text.replace(target_text, replacement_text)
                        modified = True
                        break
                else:
                    # Otherwise we replace in paragraph text
                    paragraph.text = paragraph.text.replace(target_text, replacement_text)
                    modified = True
        
        # 2. Search in tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        clean_para_text = " ".join(paragraph.text.strip().split())
                        if clean_target in clean_para_text or target_text in paragraph.text:
                            for run in paragraph.runs:
                                if target_text in run.text:
                                    run.text = run.text.replace(target_text, replacement_text)
                                    modified = True
                                    break
                            else:
                                paragraph.text = paragraph.text.replace(target_text, replacement_text)
                                modified = True
                                
        return modified

    def compact_docx_layout(self, doc, font_size_modifier=-1.0, spacing_modifier=-2.0):
        """
        Shrinks page margins, paragraph spacing, and line spacing of an existing docx Document 
        to fit content onto a single page without deleting any text content.
        """
        from docx.shared import Inches, Pt
        from docx.enum.table import WD_ROW_HEIGHT_RULE
        
        # 1. Reduce margins: adjust margin based on spacing modifier (normally 0.4 inches)
        margin_size = max(0.25, 0.4 + spacing_modifier * 0.05)
        for section in doc.sections:
            section.top_margin = Inches(margin_size)
            section.bottom_margin = Inches(margin_size)
            section.left_margin = Inches(margin_size)
            section.right_margin = Inches(margin_size)
            
        # 2. Iterate over document styles to reduce default font sizes proportionally
        for style in doc.styles:
            if hasattr(style, 'font') and style.font and style.font.size:
                try:
                    style.font.size = Pt(max(8.0, style.font.size.pt + font_size_modifier))
                except Exception:
                    pass

        # 3. Iterate over paragraphs to adjust spacing and explicit run font sizes
        for paragraph in doc.paragraphs:
            p_format = paragraph.paragraph_format
            
            # Reduce paragraph spacing
            p_format.space_before = Pt(max(1.0, 4.0 + spacing_modifier))
            p_format.space_after = Pt(max(1.0, 3.0 + spacing_modifier))
            
            # Reduce line spacing to single / compact (1.05 + adjustment)
            p_format.line_spacing = max(0.9, 1.05 + spacing_modifier * 0.02)
            
            # Reduce run font sizes slightly ONLY if explicitly set
            for run in paragraph.runs:
                if run.font.size:
                    try:
                        run.font.size = Pt(max(8.0, run.font.size.pt + font_size_modifier))
                    except Exception:
                        pass

        # 4. Iterate over tables to adjust spacing and cells
        for table in doc.tables:
            for row in table.rows:
                # Set smaller height if needed or allow row break across pages
                row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
                for cell in row.cells:
                    # Reduce paragraph spacing inside table cells
                    for paragraph in cell.paragraphs:
                        p_format = paragraph.paragraph_format
                        p_format.space_before = Pt(max(0.5, 1.0 + spacing_modifier * 0.5))
                        p_format.space_after = Pt(max(0.5, 1.0 + spacing_modifier * 0.5))
                        p_format.line_spacing = 1.0
                        for run in paragraph.runs:
                            if run.font.size:
                                run.font.size = Pt(max(7.5, run.font.size.pt + font_size_modifier))
                            else:
                                run.font.size = Pt(max(7.5, 8.5 + font_size_modifier))
                                
        return doc

    def convert_docx_to_pdf(self, docx_bytes):
        """
        Convert DOCX bytes to PDF bytes using docx2pdf (requires MS Word on Windows).
        """
        import tempfile
        import os
        from docx2pdf import convert
        import pythoncom
        
        pythoncom.CoInitialize()
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                docx_path = os.path.join(temp_dir, "temp_resume.docx")
                pdf_path = os.path.join(temp_dir, "temp_resume.pdf")
                
                with open(docx_path, "wb") as f:
                    f.write(docx_bytes)
                    
                # Convert DOCX to PDF
                convert(docx_path, pdf_path)
                
                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f:
                        pdf_data = f.read()
                    return pdf_data
                else:
                    raise FileNotFoundError("PDF conversion failed: output file not created.")
        finally:
            pythoncom.CoUninitialize()



