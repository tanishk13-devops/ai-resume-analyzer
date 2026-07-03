import streamlit as st
import tempfile
import os
import io
import re
import docx
import json
from datetime import datetime
from dotenv import load_dotenv
from modules.resume_analyzer import ResumeAnalyzer
from utils.file_handler import FileHandler
from utils.template_renderer import TemplateRenderer
from utils.db_handler import DBHandler

# Load environment variables
load_dotenv()

# Page configuration
st.set_page_config(
    page_title="JobSuite AI - Ultimate Career Suite",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for Premium Dark UI with Glassmorphism
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;800&family=JetBrains+Mono:wght@300;400;700&display=swap');

    /* Global styles */
    .stApp {
        font-family: 'Outfit', sans-serif;
        background-color: #0b0f19;
        color: #e2e8f0;
    }
    
    /* Title gradient */
    .title-gradient {
        background: linear-gradient(135deg, #60a5fa 0%, #a78bfa 50%, #f472b6 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 800;
        font-size: 2.5rem;
        margin-bottom: 5px;
    }
    
    .subtitle {
        color: #94a3b8;
        font-size: 1.05rem;
        margin-bottom: 25px;
    }
    
    /* Metrics Scorecard */
    .metric-card {
        background: rgba(15, 23, 42, 0.6);
        backdrop-filter: blur(12px);
        border: 1px solid #1e293b;
        border-radius: 16px;
        padding: 20px;
        text-align: center;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.2);
        transition: all 0.3s ease;
    }
    
    .metric-card:hover {
        transform: translateY(-2px);
        border-color: #3b82f6;
        box-shadow: 0 10px 30px rgba(59, 130, 246, 0.15);
    }
    
    .metric-value {
        font-size: 3rem;
        font-weight: 800;
        background: linear-gradient(135deg, #3b82f6, #8b5cf6);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        line-height: 1;
    }
    
    .metric-label {
        font-size: 0.85rem;
        color: #94a3b8;
        text-transform: uppercase;
        letter-spacing: 1.5px;
        margin-top: 10px;
    }
    
    /* Suggestion Cards */
    .suggestion-card {
        background: rgba(30, 41, 59, 0.4);
        border-left: 5px solid #8b5cf6;
        border-top: 1px solid #1e293b;
        border-right: 1px solid #1e293b;
        border-bottom: 1px solid #1e293b;
        border-radius: 12px;
        padding: 18px;
        margin-bottom: 16px;
        transition: all 0.2s ease;
    }
    
    .suggestion-card:hover {
        background: rgba(30, 41, 59, 0.7);
        border-left-color: #ec4899;
    }
    
    .suggestion-title {
        font-weight: 600;
        color: #f1f5f9;
        font-size: 1rem;
        margin-bottom: 6px;
    }
    
    .suggestion-impact {
        background: rgba(139, 92, 246, 0.2);
        color: #c084fc;
        font-size: 11px;
        font-weight: 600;
        padding: 2px 8px;
        border-radius: 20px;
        border: 1px solid rgba(139, 92, 246, 0.3);
        display: inline-block;
        margin-bottom: 10px;
    }
    
    .diff-box {
        background: #090d16;
        border: 1px solid #1e293b;
        border-radius: 8px;
        padding: 10px;
        font-family: 'JetBrains Mono', monospace;
        font-size: 11px;
        color: #cbd5e1;
        margin-bottom: 10px;
        overflow-x: auto;
    }
    
    .diff-old {
        color: #f87171;
        text-decoration: line-through;
        margin-bottom: 4px;
    }
    
    .diff-new {
        color: #4ade80;
    }
    
    /* Workspace resume container */
    .resume-workspace {
        font-family: 'JetBrains Mono', monospace;
        background: #090d16;
        border: 1px solid #1e293b;
        border-radius: 12px;
        padding: 20px;
        height: 500px;
        overflow-y: auto;
        white-space: pre-wrap;
        font-size: 13px;
        line-height: 1.5;
        color: #e2e8f0;
    }
    
    /* Kanban Boards */
    .kanban-col {
        background: rgba(15, 23, 42, 0.4);
        border: 1px solid #1e293b;
        border-radius: 12px;
        padding: 15px;
        min-height: 400px;
    }
    
    .kanban-card {
        background: rgba(30, 41, 59, 0.6);
        border: 1px solid #334155;
        border-radius: 10px;
        padding: 12px;
        margin-bottom: 12px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        transition: all 0.2s ease;
    }
    
    .kanban-card:hover {
        border-color: #3b82f6;
        transform: translateY(-2px);
    }
    
    .kanban-status-badge {
        font-size: 10px;
        font-weight: 700;
        text-transform: uppercase;
        padding: 2px 8px;
        border-radius: 12px;
        display: inline-block;
        margin-bottom: 6px;
    }
    
    .status-wishlist { background: rgba(234, 179, 8, 0.2); color: #facc15; }
    .status-applied { background: rgba(59, 130, 246, 0.2); color: #60a5fa; }
    .status-interviewing { background: rgba(168, 85, 247, 0.2); color: #c084fc; }
    .status-offer { background: rgba(34, 197, 94, 0.2); color: #4ade80; }
    .status-rejected { background: rgba(239, 68, 68, 0.2); color: #f87171; }
    
    /* Chat bubbles */
    .chat-bubble-assistant {
        background: rgba(30, 41, 59, 0.6);
        border: 1px solid #1e293b;
        border-radius: 12px;
        padding: 12px;
        margin-bottom: 12px;
        color: #cbd5e1;
    }
    
    .chat-bubble-user {
        background: rgba(59, 130, 246, 0.15);
        border: 1px solid rgba(59, 130, 246, 0.3);
        border-radius: 12px;
        padding: 12px;
        margin-bottom: 12px;
        color: #f1f5f9;
        text-align: right;
    }
    
    .check-item {
        display: flex;
        align-items: flex-start;
        padding: 10px 0;
        border-bottom: 1px solid #1e293b;
    }
    
    .check-icon {
        margin-right: 12px;
        font-size: 1.2rem;
    }
    
    .check-details {
        flex: 1;
    }
    
    .check-title {
        font-weight: 600;
        font-size: 0.95rem;
    }
    
    .check-desc {
        font-size: 0.85rem;
        color: #64748b;
    }

    .check-impact {
        color: #f87171;
        font-size: 0.8rem;
        font-weight: 600;
        margin-left: 10px;
    }

    /* Tabs decoration */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        background-color: rgba(15, 23, 42, 0.4);
        border: 1px solid #1e293b;
        border-radius: 8px 8px 0 0;
        padding: 10px 20px;
        color: #94a3b8;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: #1e293b !important;
        color: #3b82f6 !important;
        border-color: #3b82f6 !important;
    }
    </style>
""", unsafe_allow_html=True)

# Initialize Session States
st.session_state.analyzer = ResumeAnalyzer()
st.session_state.file_handler = FileHandler()
st.session_state.template_renderer = TemplateRenderer()

# Active workspace details
if 'active_resume_text' not in st.session_state:
    st.session_state.active_resume_text = ""
if 'structured_json' not in st.session_state:
    st.session_state.structured_json = None
if 'original_docx_doc' not in st.session_state:
    st.session_state.original_docx_doc = None
if 'original_docx_bytes' not in st.session_state:
    st.session_state.original_docx_bytes = None
if 'resume_filename' not in st.session_state:
    st.session_state.resume_filename = ""
if 'resume_format' not in st.session_state:
    st.session_state.resume_format = ""

# Analysis state
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = None
if 'applied_suggestions' not in st.session_state:
    st.session_state.applied_suggestions = set()
if 'score_increase' not in st.session_state:
    st.session_state.score_increase = 0

# Cover Letter state
if 'generated_cover_letter' not in st.session_state:
    st.session_state.generated_cover_letter = ""
if 'cover_letter_chat' not in st.session_state:
    st.session_state.cover_letter_chat = []

# Mock Interview state
if 'mock_chat' not in st.session_state:
    st.session_state.mock_chat = []
if 'mock_last_question' not in st.session_state:
    st.session_state.mock_last_question = ""
if 'interview_prep_data' not in st.session_state:
    st.session_state.interview_prep_data = None

# Helper functions
def compile_json_to_text(data):
    """Compile structured JSON data into readable plain text layout"""
    out = []
    name = data.get("name", "Resume Owner")
    out.append(name.upper())
    
    contact = data.get("contact", {})
    contact_parts = []
    if contact.get("email"): contact_parts.append(contact["email"])
    if contact.get("phone"): contact_parts.append(contact["phone"])
    if contact.get("linkedin"): contact_parts.append(contact["linkedin"])
    if contact.get("location"): contact_parts.append(contact["location"])
    out.append(" | ".join(contact_parts))
    out.append("=" * 60)
    out.append("")
    
    summary = data.get("summary")
    if summary:
        out.append("SUMMARY")
        out.append("-" * 30)
        out.append(summary)
        out.append("")
        
    experience = data.get("experience", [])
    if experience:
        out.append("PROFESSIONAL EXPERIENCE")
        out.append("-" * 30)
        for exp in experience:
            out.append(f"{exp.get('job_title', '')} | {exp.get('company', '')} ({exp.get('dates', '')})")
            for bullet in exp.get("description", []):
                if bullet.strip():
                    out.append(f"  - {bullet.strip()}")
            out.append("")
            
    skills = data.get("skills", [])
    if skills:
        out.append("TECHNICAL SKILLS")
        out.append("-" * 30)
        out.append(", ".join(skills) if isinstance(skills, list) else skills)
        out.append("")
        
    education = data.get("education", [])
    if education:
        out.append("EDUCATION")
        out.append("-" * 30)
        for edu in education:
            out.append(f"{edu.get('degree', '')} | {edu.get('school', '')} ({edu.get('dates', '')})")
            out.append("")
            
    return "\n".join(out)

# Callbacks
def apply_suggestion_callback(suggestion):
    id_ = suggestion['id']
    st.session_state.applied_suggestions.add(id_)
    
    target = suggestion.get('target_text')
    replacement = suggestion.get('replacement_text')
    st.session_state.score_increase += suggestion.get('impact_points', 0)
    
    if suggestion['type'] == 'rewrite_bullet' and target:
        if target in st.session_state.active_resume_text:
            st.session_state.active_resume_text = st.session_state.active_resume_text.replace(target, replacement)
            if st.session_state.original_docx_doc:
                st.session_state.file_handler.replace_text_in_docx(
                    st.session_state.original_docx_doc, target, replacement
                )
            st.toast(f"✅ Bullet point updated: \"{suggestion['suggestion']}\"")
        else:
            # Fuzzy match fallback
            clean_target = "".join(target.split())
            cleaned_resume = "".join(st.session_state.active_resume_text.split())
            if clean_target in cleaned_resume:
                sentences = re.split(r'(?<=[.!?])\s+', st.session_state.active_resume_text)
                for i, s in enumerate(sentences):
                    if "".join(target.split()[:5]) in "".join(s.split()):
                        st.session_state.active_resume_text = st.session_state.active_resume_text.replace(s, replacement)
                        if st.session_state.original_docx_doc:
                            st.session_state.file_handler.replace_text_in_docx(st.session_state.original_docx_doc, s, replacement)
                        break
                st.toast(f"✅ Bullet point updated: \"{suggestion['suggestion']}\"")
            else:
                st.warning("Could not find exact text in resume to replace, but active preview updated.")
                
    elif suggestion['type'] == 'add_skill':
        skills_keywords = ['skills', 'technologies', 'technical skills', 'competencies']
        found = False
        for kw in skills_keywords:
            pattern = re.compile(rf'\b{kw}\b', re.IGNORECASE)
            match = pattern.search(st.session_state.active_resume_text)
            if match:
                end_pos = match.end()
                st.session_state.active_resume_text = (
                    st.session_state.active_resume_text[:end_pos] + f"\n- {replacement}," + st.session_state.active_resume_text[end_pos:]
                )
                found = True
                break
        
        if not found:
            st.session_state.active_resume_text += f"\n\nSkills: {replacement}"
            
        if st.session_state.original_docx_doc:
            docx_found = False
            for p in st.session_state.original_docx_doc.paragraphs:
                if any(kw in p.text.lower() for kw in skills_keywords):
                    p.text = p.text + ", " + replacement
                    docx_found = True
                    break
            if not docx_found:
                st.session_state.original_docx_doc.add_paragraph(f"Skills: {replacement}")
                
        st.toast(f"➕ Added Skill: {replacement}")
        
        # Adjust local matching skills list
        if st.session_state.analysis_results:
            if 'matching_skills' in st.session_state.analysis_results:
                if replacement not in st.session_state.analysis_results['matching_skills']:
                    st.session_state.analysis_results['matching_skills'].append(replacement)
            if 'missing_skills' in st.session_state.analysis_results:
                if replacement in st.session_state.analysis_results['missing_skills']:
                    st.session_state.analysis_results['missing_skills'].remove(replacement)
                
    elif suggestion['type'] == 'add_section':
        section_name = suggestion['section']
        st.session_state.active_resume_text += f"\n\n{section_name.upper()}\n- {replacement}"
        if st.session_state.original_docx_doc:
            st.session_state.original_docx_doc.add_heading(section_name.title(), level=2)
            st.session_state.original_docx_doc.add_paragraph(replacement)
        st.toast(f"➕ Added Section: {section_name.title()}")

def add_header_callback(info_type, value):
    if not value.strip():
        return
    header = f"{info_type.capitalize()}: {value}"
    st.session_state.active_resume_text = header + "\n" + st.session_state.active_resume_text
    
    if st.session_state.original_docx_doc:
        p = st.session_state.original_docx_doc.insert_paragraph_before(header)
        if len(st.session_state.original_docx_doc.paragraphs) > 1:
            p.style = st.session_state.original_docx_doc.paragraphs[1].style
            
    st.session_state.score_increase += 10
    st.toast(f"✅ Added {info_type.capitalize()}: {value}")

# Sidebar Navigation Header
st.sidebar.markdown("""
    <div style="text-align: center; margin-bottom: 20px;">
        <h1 style="color: #60a5fa; font-weight: 800; font-size: 1.8rem; margin: 0;">JobSuite AI</h1>
        <p style="color: #94a3b8; font-size: 0.85rem; margin-top: 5px;">Your Complete AI Career Assistant</p>
    </div>
""", unsafe_allow_html=True)

nav_page = st.sidebar.radio(
    "🧭 Navigate Suite Modules:",
    [
        "🎯 Resume Optimizer & Tailoring",
        "🛠️ Interactive Resume Builder",
        "✉️ AI Cover Letter Generator",
        "📊 Job Board CRM Tracker",
        "🧠 Interview Prep & Mock Chat"
    ]
)

st.sidebar.markdown("---")

# Global LLM Model Engine Selector
st.sidebar.markdown("### 🤖 LLM Model Engine")
is_gemini = os.getenv('GEMINI_API_KEY') and (not os.getenv('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') == "your_openai_key_here")

if is_gemini:
    model_options = [
        "gemini-2.0-flash-lite",
        "gemini-2.0-flash",
        "gemini-2.5-flash",
        "gemini-2.5-flash-lite",
        "gemini-flash-latest",
        "gemini-pro-latest"
    ]
else:
    model_options = [
        "gpt-3.5-turbo",
        "gpt-4o-mini",
        "gpt-4"
    ]
    
selected_model = st.sidebar.selectbox(
    "Select API Model:",
    model_options,
    index=0,
    key="selected_llm_model"
)

# Apply model to handler
if 'analyzer' in st.session_state:
    st.session_state.analyzer.llm_handler.model = selected_model

st.sidebar.markdown("---")


# ==========================================
# MODULE 1: RESUME OPTIMIZER & TAILORING
# ==========================================
if "Resume Optimizer & Tailoring" in nav_page:
    st.markdown('<div class="title-gradient">🎯 Resume Optimizer & Tailoring</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitle">Tailor and match your resume to specific target jobs, run ATS checks, and apply quick upgrades</div>', unsafe_allow_html=True)
    
    # Configuration options in the sidebar for this page
    with st.sidebar:
        st.markdown("### ⚙️ Optimization Config")
        
        # Upload option
        resume_source = st.radio(
            "Select Resume Source:",
            ["Upload File", "Import from Resume Builder"],
            index=0 if st.session_state.active_resume_text == "" else 1
        )
        
        uploaded_resume = None
        if resume_source == "Upload File":
            uploaded_resume = st.file_uploader(
                "Upload Resume (PDF/DOCX)",
                type=["pdf", "docx"],
                key="opt_file_uploader"
            )
            if uploaded_resume:
                if st.session_state.resume_filename != uploaded_resume.name:
                    st.session_state.resume_filename = uploaded_resume.name
                    st.session_state.resume_format = uploaded_resume.name.split('.')[-1].lower()
                    st.session_state.applied_suggestions = set()
                    st.session_state.score_increase = 0
                    
                    if st.session_state.resume_format == 'docx':
                        st.session_state.original_docx_bytes = uploaded_resume.getvalue()
                        st.session_state.original_docx_doc = docx.Document(io.BytesIO(st.session_state.original_docx_bytes))
                        text = ""
                        for p in st.session_state.original_docx_doc.paragraphs:
                            text += p.text + "\n"
                        st.session_state.active_resume_text = text
                    else:
                        st.session_state.original_docx_doc = None
                        st.session_state.original_docx_bytes = None
                        st.session_state.active_resume_text = st.session_state.file_handler.extract_text(uploaded_resume)
                    
                    # Reset analysis
                    st.session_state.analysis_results = None
        else:
            if st.session_state.structured_json:
                st.markdown("✅ *Resume Builder data loaded!*")
            else:
                st.info("⚠️ No Resume Builder data found. Create a resume in the Builder first!")
                
        st.markdown("---")
        st.markdown("### 📋 Target Job Description")
        job_input_method = st.radio("Input Job Description:", ["Paste Text", "Upload File"], key="opt_job_method")
        
        if job_input_method == "Paste Text":
            job_description = st.text_area("Paste job requirements:", height=180, placeholder="Paste job description here...")
            job_file = None
        else:
            job_file = st.file_uploader("Upload Job Description (PDF/DOCX)", type=["pdf", "docx"])
            job_description = None
            
        st.markdown("---")
        analyze_btn = st.button("🚀 Analyze & Optimize", use_container_width=True, key="opt_run_btn")

    # Business logic
    if (st.session_state.active_resume_text or uploaded_resume) and (job_description or job_file):
        if analyze_btn or st.session_state.analysis_results is None:
            try:
                with st.spinner("Analyzing resume content and matching keyword database..."):
                    if job_file:
                        job_text = st.session_state.file_handler.extract_text(job_file)
                    else:
                        job_text = job_description
                        
                    # Save target job description in session state for other tabs
                    st.session_state.target_job_desc = job_text
                    
                    st.session_state.score_increase = 0
                    st.session_state.applied_suggestions = set()
                    
                    st.session_state.analysis_results = st.session_state.analyzer.run_full_analysis(
                        st.session_state.active_resume_text,
                        job_text
                    )
                    st.toast("🎉 Analysis complete!")
            except Exception as e:
                st.error(f"Error running analysis: {str(e)}")
                st.session_state.analysis_results = None

        if st.session_state.analysis_results:
            results = st.session_state.analysis_results
            
            raw_match = float(results.get("match_score", 0))
            raw_skills = float(results.get("skills_match", 0))
            formatting_data = st.session_state.analyzer.text_processor.check_ats_formatting(
                st.session_state.active_resume_text
            )
            formatting_score = formatting_data["score"]
            adjusted_match = min(100.0, raw_match + st.session_state.score_increase)
            combined_ats_score = (adjusted_match * 0.6) + (formatting_score * 0.4)
            
            # Scores row
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown(f"""
                    <div class="metric-card">
                        <div class="metric-value">{combined_ats_score:.0f}%</div>
                        <div class="metric-label">Combined ATS Score</div>
                    </div>
                """, unsafe_allow_html=True)
            with col2:
                st.markdown(f"""
                    <div class="metric-card">
                        <div class="metric-value">{adjusted_match:.0f}%</div>
                        <div class="metric-label">Content Match Score</div>
                    </div>
                """, unsafe_allow_html=True)
            with col3:
                st.markdown(f"""
                    <div class="metric-card">
                        <div class="metric-value">{formatting_score:.0f}%</div>
                        <div class="metric-label">Layout Formatting Score</div>
                    </div>
                """, unsafe_allow_html=True)
                
            st.markdown("<br>", unsafe_allow_html=True)
            
            tab1, tab2, tab3 = st.tabs([
                "🛠️ Interactive Workspace", 
                "📊 Full Analysis Report", 
                "📥 Export & Download"
            ])
            
            with tab1:
                workspace_col1, workspace_col2 = st.columns([1.1, 0.9])
                with workspace_col1:
                    st.markdown("**Active Resume Preview** (Updated dynamically)")
                    st.markdown(f'<div class="resume-workspace">{st.session_state.active_resume_text}</div>', unsafe_allow_html=True)
                    
                    with st.expander("📝 Manual Editor Mode"):
                        manual_text = st.text_area(
                            "Tweak resume wording directly:",
                            value=st.session_state.active_resume_text,
                            height=250
                        )
                        if manual_text != st.session_state.active_resume_text:
                            st.session_state.active_resume_text = manual_text
                            st.experimental_rerun()
                
                with workspace_col2:
                    st.markdown("**Instant AI & Layout Upgrades**")
                    
                    # Heuristics updates
                    failed_checks = formatting_data.get("failed", [])
                    if failed_checks:
                        st.markdown("##### ⚠️ Layout Improvements Required")
                        for fc in failed_checks:
                            with st.container():
                                st.markdown(f"""
                                    <div style="background: rgba(248, 113, 113, 0.05); border: 1px solid rgba(248, 113, 113, 0.2); border-radius: 8px; padding: 12px; margin-bottom: 10px;">
                                        <strong style="color: #f87171;">{fc['check']}</strong>
                                        <div style="font-size: 0.85rem; color: #94a3b8; margin: 4px 0 8px 0;">{fc['tip']}</div>
                                    </div>
                                """, unsafe_allow_html=True)
                                
                                if "Email" in fc['check']:
                                    email_val = st.text_input("Add Email Address:", placeholder="you@domain.com", key="opt_inline_email")
                                    if st.button("➕ Add Email", key="opt_btn_email"):
                                        add_header_callback('email', email_val)
                                        st.experimental_rerun()
                                elif "Phone" in fc['check']:
                                    phone_val = st.text_input("Add Phone Number:", placeholder="+1 (555) 123-4567", key="opt_inline_phone")
                                    if st.button("➕ Add Phone", key="opt_btn_phone"):
                                        add_header_callback('phone', phone_val)
                                        st.experimental_rerun()
                                elif "LinkedIn" in fc['check']:
                                    li_val = st.text_input("Add LinkedIn Link:", placeholder="linkedin.com/in/username", key="opt_inline_li")
                                    if st.button("➕ Add LinkedIn", key="opt_btn_li"):
                                        add_header_callback('linkedin', li_val)
                                        st.experimental_rerun()
                                        
                    # AI Content suggestions
                    suggestions = results.get("suggestions", [])
                    visible = [s for s in suggestions if s['id'] not in st.session_state.applied_suggestions]
                    
                    if visible:
                        st.markdown("##### ✨ AI Content Enhancements")
                        for sug in visible:
                            with st.container():
                                st.markdown(f"""
                                    <div class="suggestion-card">
                                        <div class="suggestion-title">{sug['suggestion']}</div>
                                        <div class="suggestion-impact">+{sug['impact_points']}% Match Score</div>
                                    </div>
                                """, unsafe_allow_html=True)
                                
                                if sug['type'] == 'rewrite_bullet' and sug.get('target_text'):
                                    st.markdown(f"""
                                        <div class="diff-box">
                                            <div class="diff-old">- {sug['target_text']}</div>
                                            <div class="diff-new">+ {sug['replacement_text']}</div>
                                        </div>
                                    """, unsafe_allow_html=True)
                                elif sug['type'] == 'add_skill':
                                    st.markdown(f"""
                                        <div class="diff-box">
                                            <div class="diff-new">➕ Add Skill: {sug['replacement_text']}</div>
                                        </div>
                                    """, unsafe_allow_html=True)
                                    
                                st.button(
                                    "⚡ Apply Upgrade",
                                    key=f"opt_apply_{sug['id']}",
                                    on_click=apply_suggestion_callback,
                                    args=(sug,)
                                )
                    else:
                        st.success("🎉 All resume tailoring suggestions applied!")
                        
                    if st.button("🔄 Reset Optimization Adjustments", key="opt_reset_btn"):
                        st.session_state.applied_suggestions = set()
                        st.session_state.score_increase = 0
                        if st.session_state.resume_format == 'docx' and st.session_state.original_docx_bytes:
                            st.session_state.original_docx_doc = docx.Document(io.BytesIO(st.session_state.original_docx_bytes))
                            text = ""
                            for p in st.session_state.original_docx_doc.paragraphs:
                                text += p.text + "\n"
                            st.session_state.active_resume_text = text
                        elif uploaded_resume:
                            st.session_state.active_resume_text = st.session_state.file_handler.extract_text(uploaded_resume)
                        elif st.session_state.structured_json:
                            st.session_state.active_resume_text = compile_json_to_text(st.session_state.structured_json)
                        st.experimental_rerun()
                        
            with tab2:
                rep_col1, rep_col2 = st.columns(2)
                with rep_col1:
                    st.subheader("💡 Key Insights")
                    st.markdown("**Top Strengths:**")
                    for s in results.get("strengths", []):
                        st.write(f"✨ {s}")
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.markdown("**Potential Concerns:**")
                    for c in results.get("concerns", []):
                        st.write(f"⚠️ {c}")
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.markdown(f"**Education Assessment:**\n{results.get('education_assessment', 'N/A')}")
                    
                with rep_col2:
                    st.subheader("🔑 Keywords Matching")
                    st.metric("Keywords & Skills Match Rate", f"{raw_skills:.0f}%")
                    
                    sk_col1, sk_col2 = st.columns(2)
                    with sk_col1:
                        st.markdown("**✅ Matching Skills**")
                        for sk in results.get("matching_skills", [])[:8]:
                            st.success(sk)
                    with sk_col2:
                        st.markdown("**❌ Missing Skills**")
                        for sk in results.get("missing_skills", [])[:8]:
                            st.warning(sk)
                            
            with tab3:
                st.markdown("### Export & Download Options")
                exp_col1, exp_col2 = st.columns([1, 1.2])
                with exp_col1:
                    st.markdown("##### 1. Select Export Layout Structure")
                    
                    export_options = []
                    if st.session_state.resume_format == 'docx':
                        export_options.append("Keep Original Uploaded Format (DOCX / TXT)")
                    else:
                        export_options.append("Keep Original Uploaded Format (TXT Only)")
                    export_options.append("Convert to Selected Layout Template (PDF / DOCX / TXT)")
                    
                    exp_mode = st.radio(
                        "Choose layout mode:",
                        export_options,
                        index=0,
                        key="opt_exp_layout_mode"
                    )
                    
                    style_key = "executive_classic"
                    if "Convert to Selected Layout" in exp_mode:
                        template_options = st.session_state.template_renderer.get_template_keys()
                        template_display_names = [name for key, name in template_options]
                        template_keys = [key for key, name in template_options]
                        
                        selected_name = st.selectbox(
                            "Select ATS-Friendly Style Template (25 options):",
                            template_display_names,
                            key="opt_template_select"
                        )
                        style_key = template_keys[template_display_names.index(selected_name)]
                        st.success(f"Selected: {selected_name}")
                        
                        # Fit to Single Page toggle
                        fit_single = st.checkbox(
                            "Fit to Single Page (Compact Spacing & Margins)",
                            value=False,
                            key="opt_fit_single_page"
                        )
                        
                        font_mod = 0.0
                        space_mod = 0.0
                        if fit_single:
                            st.caption("⚙️ Fine-tune Fit Spacing:")
                            font_mod = st.slider("Font Size Adjust (pt):", -2.5, 1.0, -1.0, 0.1, key="opt_fit_font_mod")
                            space_mod = st.slider("Spacing Adjust (pt):", -4.0, 2.0, -2.0, 0.2, key="opt_fit_space_mod")
                    else:
                        st.info("💡 Keeping original formatting coordinates. Direct text updates will be injected into your original layout structure.")
                        
                with exp_col2:
                    st.markdown("##### 2. File Download Actions")
                    
                    if "Convert to Selected Layout" in exp_mode:
                        # Lazy parsing to json database if missing (only if not imported from Resume Builder)
                        if resume_source != "Import from Resume Builder" and ('structured_json' not in st.session_state or st.session_state.get('last_parsed_text') != st.session_state.active_resume_text):
                            with st.spinner("Compiling structured resume database for templates..."):
                                st.session_state.structured_json = st.session_state.analyzer.llm_handler.parse_resume_to_json(
                                    st.session_state.active_resume_text
                                )
                                st.session_state.last_parsed_text = st.session_state.active_resume_text
                                if st.session_state.structured_json and st.session_state.structured_json.get("name") == "Resume Owner":
                                    st.warning("⚠️ The selected Gemini model is currently overloaded or rate-limited (503). The app fell back to a basic local parser. Try selecting a different model in the sidebar (e.g. gemini-2.0-flash) and re-running.")
                        
                        struct = st.session_state.structured_json
                        pdf_bytes = st.session_state.template_renderer.generate_pdf(
                            struct, style_key, fit_single_page=fit_single, font_size_modifier=font_mod, spacing_modifier=space_mod
                        )
                        docx_bytes = st.session_state.template_renderer.generate_docx(
                            struct, style_key, fit_single_page=fit_single, font_size_modifier=font_mod, spacing_modifier=space_mod
                        )
                        txt_bytes = st.session_state.template_renderer.generate_txt(struct).encode('utf-8')
                        
                        st.download_button(
                            "📥 Download Resume as PDF (Template)",
                            data=pdf_bytes,
                            file_name=f"jobsuite_resume_{style_key}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                        st.download_button(
                            "📥 Download Resume as DOCX (Template)",
                            data=docx_bytes,
                            file_name=f"jobsuite_resume_{style_key}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                        st.download_button(
                            "📥 Download Resume as TXT (Template)",
                            data=txt_bytes,
                            file_name=f"jobsuite_resume_{style_key}.txt",
                            mime="text/plain",
                            use_container_width=True
                        )
                    else:
                        fit_single_orig = st.checkbox(
                            "Fit to Single Page (Compact Spacing & Margins)",
                            value=False,
                            key="opt_fit_single_original"
                        )
                        
                        font_mod_orig = 0.0
                        space_mod_orig = 0.0
                        if fit_single_orig:
                            st.caption("⚙️ Fine-tune Fit Spacing:")
                            font_mod_orig = st.slider("Font Size Adjust (pt):", -2.5, 1.0, -1.0, 0.1, key="opt_fit_font_mod_orig")
                            space_mod_orig = st.slider("Spacing Adjust (pt):", -4.0, 2.0, -2.0, 0.2, key="opt_fit_space_mod_orig")
                        
                        docx_bytes_out = None
                        pdf_bytes_out = None
                        
                        # A. Uploaded file was a DOCX
                        if st.session_state.resume_format == 'docx' and st.session_state.original_docx_doc:
                            # 1. Generate DOCX
                            bio_temp = io.BytesIO()
                            st.session_state.original_docx_doc.save(bio_temp)
                            bio_temp.seek(0)
                            doc_to_save = docx.Document(bio_temp)
                            
                            if fit_single_orig:
                                doc_to_save = st.session_state.file_handler.compact_docx_layout(
                                    doc_to_save, font_size_modifier=font_mod_orig, spacing_modifier=space_mod_orig
                                )
                                
                            bio = io.BytesIO()
                            doc_to_save.save(bio)
                            docx_bytes_out = bio.getvalue()
                            
                            # 2. Generate PDF
                            try:
                                with st.spinner("Converting original layout to PDF..."):
                                    pdf_bytes_out = st.session_state.file_handler.convert_docx_to_pdf(docx_bytes_out)
                            except Exception as e:
                                # Fallback to Traditional/Classic template PDF if COM fails or not installed
                                if resume_source != "Import from Resume Builder" and not st.session_state.structured_json:
                                    st.session_state.structured_json = st.session_state.analyzer.llm_handler.parse_resume_to_json(
                                        st.session_state.active_resume_text
                                    )
                                    if st.session_state.structured_json and st.session_state.structured_json.get("name") == "Resume Owner":
                                        st.warning("⚠️ Gemini model overloaded/rate-limited. Local parsing fallback used.")
                                pdf_bytes_out = st.session_state.template_renderer.generate_pdf(
                                    st.session_state.structured_json, 
                                    template_id="executive_classic", 
                                    fit_single_page=fit_single_orig,
                                    font_size_modifier=font_mod_orig,
                                    spacing_modifier=space_mod_orig
                                )
                                st.warning("⚠️ Word COM automation is unavailable. PDF was successfully rendered using the Classic Template.")
                        
                        # B. Uploaded file was a PDF or other format (render standard layout as the closest layout representation)
                        else:
                            if resume_source != "Import from Resume Builder" and not st.session_state.structured_json:
                                st.session_state.structured_json = st.session_state.analyzer.llm_handler.parse_resume_to_json(
                                    st.session_state.active_resume_text
                                )
                                if st.session_state.structured_json and st.session_state.structured_json.get("name") == "Resume Owner":
                                    st.warning("⚠️ The selected Gemini model is currently overloaded or rate-limited (503). The app fell back to a basic local parser. Try selecting a different model in the sidebar.")
                            
                            pdf_bytes_out = st.session_state.template_renderer.generate_pdf(
                                st.session_state.structured_json, 
                                template_id="executive_classic", 
                                fit_single_page=fit_single_orig,
                                font_size_modifier=font_mod_orig,
                                spacing_modifier=space_mod_orig
                            )
                            docx_bytes_out = st.session_state.template_renderer.generate_docx(
                                st.session_state.structured_json, 
                                template_id="executive_classic", 
                                fit_single_page=fit_single_orig,
                                font_size_modifier=font_mod_orig,
                                spacing_modifier=space_mod_orig
                            )
                        
                        # Align buttons in 3 columns
                        col_orig1, col_orig2, col_orig3 = st.columns(3)
                        with col_orig1:
                            st.download_button(
                                "📥 Download PDF",
                                data=pdf_bytes_out,
                                file_name=f"optimized_{st.session_state.resume_filename.split('.')[0]}.pdf" if st.session_state.resume_filename else "optimized_resume.pdf",
                                mime="application/pdf",
                                use_container_width=True,
                                key="opt_dl_orig_pdf"
                            )
                        with col_orig2:
                            st.download_button(
                                "📥 Download DOCX",
                                data=docx_bytes_out,
                                file_name=f"optimized_{st.session_state.resume_filename}" if st.session_state.resume_filename else "optimized_resume.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="opt_dl_orig_docx"
                            )
                        with col_orig3:
                            st.download_button(
                                "📥 Download TXT",
                                data=st.session_state.active_resume_text.encode('utf-8'),
                                file_name="optimized_resume.txt",
                                mime="text/plain",
                                use_container_width=True,
                                key="opt_dl_orig_txt"
                            )
    else:
        st.markdown("""
            <div style="background: rgba(15, 23, 42, 0.4); border: 1px solid #1e293b; border-radius: 16px; padding: 40px; text-align: center; max-width: 800px; margin: 40px auto;">
                <div style="font-size: 4rem; margin-bottom: 20px;">👈</div>
                <h3 style="color: #f1f5f9; font-weight: 800;">Get Started in the Sidebar</h3>
                <p style="color: #94a3b8; font-size: 1rem; margin-bottom: 20px;">
                    Upload your resume (PDF or DOCX) or import data from the Interactive Resume Builder, then supply a job description to trigger dynamic scoring, keyword matching, and bullet replacements.
                </p>
            </div>
        """, unsafe_allow_html=True)


# ==========================================
# MODULE 2: INTERACTIVE RESUME BUILDER
# ==========================================
elif "Interactive Resume Builder" in nav_page:
    st.markdown('<div class="title-gradient">🛠️ Interactive Resume Builder</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitle">Build an ATS-optimized resume from scratch using dynamic step-by-step form panels and AI bullet/summary copilots</div>', unsafe_allow_html=True)
    
    # Initialize basic empty dictionary if none exists
    if not st.session_state.structured_json:
        st.session_state.structured_json = {
            "name": "",
            "contact": {"email": "", "phone": "", "linkedin": "", "location": ""},
            "summary": "",
            "experience": [],
            "skills": [],
            "education": []
        }
    
    struct = st.session_state.structured_json
    
    edit_mode = st.radio("Editor Panel:", ["Form Wizard (Recommended)", "Raw JSON Editor"], key="builder_edit_mode")
    
    if edit_mode == "Raw JSON Editor":
        raw_json_str = st.text_area("Edit resume database structure manually:", value=json.dumps(struct, indent=2), height=450)
        try:
            st.session_state.structured_json = json.loads(raw_json_str)
            st.success("JSON Database parsed correctly!")
        except Exception as e:
            st.error(f"Invalid JSON Format: {str(e)}")
            
        if st.button("💾 Compile & Push to Suite Workspace", key="builder_json_compile"):
            st.session_state.active_resume_text = compile_json_to_text(st.session_state.structured_json)
            st.toast("✅ Workspace updated successfully!")
    else:
        # Multi-tab step-by-step form wizard
        btab1, btab2, btab3, btab4, btab5 = st.tabs([
            "👤 Contact Details",
            "📝 Executive Summary",
            "💼 Professional History",
            "🛠️ Technical Skills",
            "🎓 Education History"
        ])
        
        with btab1:
            st.markdown("#### Personal & Professional Contacts")
            struct["name"] = st.text_input("Full Name:", value=struct.get("name", ""), key="b_name")
            
            col1, col2 = st.columns(2)
            with col1:
                struct["contact"]["email"] = st.text_input("Email Address:", value=struct["contact"].get("email", ""), key="b_email")
                struct["contact"]["phone"] = st.text_input("Phone Number:", value=struct["contact"].get("phone", ""), key="b_phone")
            with col2:
                struct["contact"]["linkedin"] = st.text_input("LinkedIn Profile URL:", value=struct["contact"].get("linkedin", ""), key="b_linkedin")
                struct["contact"]["location"] = st.text_input("Location (City, State/Country):", value=struct["contact"].get("location", ""), key="b_location")
                
        with btab2:
            st.markdown("#### Professional Summary")
            st.markdown("Write a compelling career introduction or let the AI write one based on your experience history.")
            
            # AI summary assist
            target_role = st.text_input("Target Job Title (for AI summary optimization):", placeholder="e.g. Senior Software Engineer", key="b_target_title")
            if st.button("✨ Draft Summary with AI", key="b_summary_ai_btn"):
                if not struct["experience"] and not struct["skills"]:
                    st.warning("Please fill out some work experience or skills first so the AI has context about your background!")
                else:
                    with st.spinner("AI drafting resume summary..."):
                        personal = {"name": struct.get("name", ""), "title": target_role}
                        drafted = st.session_state.analyzer.llm_handler.generate_summary_suggestion(
                            personal, struct["experience"], struct["skills"], target_role
                        )
                        struct["summary"] = drafted
                        st.success("Draft generated successfully!")
                        
            struct["summary"] = st.text_area("Summary content:", value=struct.get("summary", ""), height=150, key="b_summary")
            
        with btab3:
            st.markdown("#### Work Experience")
            st.markdown("Add details about your professional history. Use the AI enhancer to convert simple statements into metric-rich achievements.")
            
            # Show list of jobs
            for i, exp in enumerate(struct.get("experience", [])):
                with st.expander(f"💼 Role {i+1}: {exp.get('job_title', 'Title')} at {exp.get('company', 'Company')}", expanded=True):
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        exp["job_title"] = st.text_input(f"Job Title:", value=exp.get("job_title", ""), key=f"b_job_title_{i}")
                    with col2:
                        exp["company"] = st.text_input(f"Company / Employer:", value=exp.get("company", ""), key=f"b_company_{i}")
                    with col3:
                        exp["dates"] = st.text_input(f"Dates (e.g. 2021 - Present):", value=exp.get("dates", ""), key=f"b_dates_{i}")
                        
                    st.markdown("**Bullet Point Achievements:**")
                    bullets = exp.get("description", [])
                    
                    # Edit bullets
                    for j, bullet in enumerate(bullets):
                        bullet_col, bullet_btn_col = st.columns([8, 2.5])
                        with bullet_col:
                            bullets[j] = st.text_input(f"Achievement {j+1}:", value=bullet, key=f"b_bullet_{i}_{j}")
                        with bullet_btn_col:
                            # Bullet AI optimizer
                            if st.button("⚡ Enhance Bullet", key=f"b_bullet_ai_{i}_{j}"):
                                if not bullets[j].strip():
                                    st.warning("Please type a basic description first.")
                                else:
                                    with st.spinner("Optimizing bullet..."):
                                        enhanced = st.session_state.analyzer.llm_handler.generate_bullet_suggestion(
                                            exp["job_title"], exp["company"], bullets[j]
                                        )
                                        bullets[j] = enhanced
                                        st.experimental_rerun()
                                        
                    # Add/remove bullet controls
                    col_b1, col_b2 = st.columns(2)
                    with col_b1:
                        if st.button("➕ Add Bullet Point", key=f"b_add_bullet_{i}"):
                            bullets.append("")
                            st.experimental_rerun()
                    with col_b2:
                        if st.button("❌ Remove Bullet Point", key=f"b_rem_bullet_{i}"):
                            if bullets:
                                bullets.pop()
                                st.experimental_rerun()
                    
                    exp["description"] = bullets
                    
                    st.markdown("---")
                    if st.button("🗑️ Delete Job Role", key=f"b_del_job_{i}"):
                        struct["experience"].pop(i)
                        st.experimental_rerun()
                        
            if st.button("➕ Add Work Experience Role", key="b_add_job_btn"):
                struct["experience"].append({
                    "job_title": "",
                    "company": "",
                    "dates": "",
                    "description": [""]
                })
                st.experimental_rerun()
                
        with btab4:
            st.markdown("#### Technical Skills")
            skills_input = st.text_area("Enter your skills (separated by commas):", value=", ".join(struct.get("skills", [])), key="b_skills_raw")
            struct["skills"] = [s.strip() for s in skills_input.split(",") if s.strip()]
            
            # Show list of tags
            st.markdown("**Active Skill Tags:**")
            st.markdown(" ".join([f'<span style="background-color:#1e293b; color:#60a5fa; border:1px solid #3b82f6; padding:4px 10px; border-radius:15px; font-size:12px; margin-right:5px; margin-bottom:5px; display:inline-block;">{s}</span>' for s in struct["skills"]]), unsafe_allow_html=True)
            
        with btab5:
            st.markdown("#### Education History")
            for i, edu in enumerate(struct.get("education", [])):
                with st.expander(f"🎓 Degree {i+1}: {edu.get('degree', 'Degree')} from {edu.get('school', 'School')}", expanded=True):
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        edu["degree"] = st.text_input("Degree / Major:", value=edu.get("degree", ""), key=f"b_degree_{i}")
                    with col2:
                        edu["school"] = st.text_input("School / University:", value=edu.get("school", ""), key=f"b_school_{i}")
                    with col3:
                        edu["dates"] = st.text_input("Dates / Year of Graduation:", value=edu.get("dates", ""), key=f"b_edu_dates_{i}")
                        
                    if st.button("🗑️ Delete Education", key=f"b_del_edu_{i}"):
                        struct["education"].pop(i)
                        st.experimental_rerun()
                        
            if st.button("➕ Add Education Profile", key="b_add_edu_btn"):
                struct["education"].append({
                    "degree": "",
                    "school": "",
                    "dates": ""
                })
                st.experimental_rerun()
                
        st.markdown("---")
        
        # Save and Apply Actions
        col_act1, col_act2 = st.columns(2)
        with col_act1:
            if st.button("💾 Push to Suite Workspace & Optimize", key="b_save_workspace_btn"):
                # Compile to txt
                st.session_state.active_resume_text = compile_json_to_text(struct)
                st.session_state.structured_json = struct
                st.toast("✅ Saved to active workspace! You can optimize it or build letters now.")
        with col_act2:
            template_options = st.session_state.template_renderer.get_template_keys()
            template_display_names = [name for key, name in template_options]
            template_keys = [key for key, name in template_options]
            
            selected_name = st.selectbox(
                "Export Resume directly in Style:",
                template_display_names,
                key="builder_direct_style_select"
            )
            style_key = template_keys[template_display_names.index(selected_name)]
            
            # Fit to Single Page toggle in Builder
            fit_single_builder = st.checkbox(
                "Fit to Single Page (Compact Spacing & Margins)",
                value=False,
                key="builder_fit_single_page"
            )
            
            font_mod_b = 0.0
            space_mod_b = 0.0
            if fit_single_builder:
                st.caption("⚙️ Fine-tune Fit Spacing:")
                font_mod_b = st.slider("Font Size Adjust (pt):", -2.5, 1.0, -1.0, 0.1, key="builder_fit_font_mod")
                space_mod_b = st.slider("Spacing Adjust (pt):", -4.0, 2.0, -2.0, 0.2, key="builder_fit_space_mod")
            
            pdf_bytes = st.session_state.template_renderer.generate_pdf(
                struct, style_key, fit_single_page=fit_single_builder, font_size_modifier=font_mod_b, spacing_modifier=space_mod_b
            )
            docx_bytes = st.session_state.template_renderer.generate_docx(
                struct, style_key, fit_single_page=fit_single_builder, font_size_modifier=font_mod_b, spacing_modifier=space_mod_b
            )
            txt_bytes = st.session_state.template_renderer.generate_txt(struct).encode('utf-8')
            
            col_dl1, col_dl2, col_dl3 = st.columns(3)
            with col_dl1:
                st.download_button(
                    "📥 Download PDF",
                    data=pdf_bytes,
                    file_name=f"jobsuite_builder_resume_{style_key}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key="builder_dl_pdf"
                )
            with col_dl2:
                st.download_button(
                    "📥 Download DOCX",
                    data=docx_bytes,
                    file_name=f"jobsuite_builder_resume_{style_key}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="builder_dl_docx"
                )
            with col_dl3:
                st.download_button(
                    "📥 Download TXT",
                    data=txt_bytes,
                    file_name=f"jobsuite_builder_resume_{style_key}.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key="builder_dl_txt"
                )


# ==========================================
# MODULE 3: AI COVER LETTER GENERATOR
# ==========================================
elif "AI Cover Letter Generator" in nav_page:
    st.markdown('<div class="title-gradient">✉️ AI Cover Letter Generator</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitle">Generate tailored professional cover letters and iteratively refine them with your personal AI Career assistant</div>', unsafe_allow_html=True)
    
    # Check dependencies
    if not st.session_state.active_resume_text:
        st.warning("⚠️ No active resume found in workspace. Upload a resume in the Optimizer or build one in the Resume Builder first!")
    
    # JD setup
    jd_input = st.session_state.get('target_job_desc', "")
    
    col_input1, col_input2 = st.columns(2)
    with col_input1:
        st.markdown("##### Target Job Description")
        jd_text = st.text_area(
            "Verify target role details:", 
            value=jd_input, 
            height=150, 
            placeholder="Paste target job description requirements here...",
            key="cl_jd_text"
        )
    with col_input2:
        st.markdown("##### Generation Instructions")
        prompt_notes = st.text_area(
            "Tone requirements / custom instructions (Optional):", 
            height=150, 
            placeholder="e.g. 'Highlight my experience with Kubernetes and keep it to 300 words' or 'Write in a casual start-up tone'",
            key="cl_prompt_notes"
        )
        
    generate_cl_btn = st.button("🚀 Generate Tailored Cover Letter", use_container_width=True)
    
    if generate_cl_btn:
        if not st.session_state.active_resume_text:
            st.error("Please upload or build a resume first.")
        elif not jd_text.strip():
            st.error("Please provide the target Job Description.")
        else:
            with st.spinner("AI drafting customized cover letter matching keywords..."):
                st.session_state.generated_cover_letter = st.session_state.analyzer.llm_handler.generate_cover_letter(
                    st.session_state.active_resume_text,
                    jd_text,
                    prompt_notes
                )
                st.session_state.cover_letter_chat = [
                    {"role": "assistant", "content": "Here is the first draft of your cover letter. Let me know if you would like any specific changes!"}
                ]
                st.experimental_rerun()
                
    if st.session_state.generated_cover_letter:
        st.markdown("---")
        
        # Interactive Layout
        cl_col1, cl_col2 = st.columns([1.1, 0.9])
        
        with cl_col1:
            st.markdown("**Live Cover Letter Document** (Editable)")
            edited_cl = st.text_area(
                "Document Text Area:", 
                value=st.session_state.generated_cover_letter, 
                height=450,
                key="cl_live_editor"
            )
            if edited_cl != st.session_state.generated_cover_letter:
                st.session_state.generated_cover_letter = edited_cl
                
            # Export Options
            st.markdown("##### Match Resume Layout styling")
            
            template_options = st.session_state.template_renderer.get_template_keys()
            template_display_names = [name for key, name in template_options]
            template_keys = [key for key, name in template_options]
            
            selected_name = st.selectbox(
                "Choose Matching Template Theme:",
                template_display_names,
                key="cl_template_select"
            )
            style_key = template_keys[template_display_names.index(selected_name)]
            
            # Generate exports
            # Compile metadata for letter header
            metadata = {
                "name": "Applicant Name",
                "email": "",
                "phone": "",
                "linkedin": "",
                "location": ""
            }
            if st.session_state.structured_json:
                metadata["name"] = st.session_state.structured_json.get("name", "Applicant Name")
                metadata["email"] = st.session_state.structured_json.get("contact", {}).get("email", "")
                metadata["phone"] = st.session_state.structured_json.get("contact", {}).get("phone", "")
                metadata["linkedin"] = st.session_state.structured_json.get("contact", {}).get("linkedin", "")
                metadata["location"] = st.session_state.structured_json.get("contact", {}).get("location", "")
            
            # Exporters
            pdf_bytes = st.session_state.template_renderer.generate_cover_letter_pdf(
                st.session_state.generated_cover_letter, metadata, style_key
            )
            docx_bytes = st.session_state.template_renderer.generate_cover_letter_docx(
                st.session_state.generated_cover_letter, metadata, style_key
            )
            txt_bytes = st.session_state.generated_cover_letter.encode('utf-8')
            
            # Downloads panel
            dl_col1, dl_col2, dl_col3 = st.columns(3)
            with dl_col1:
                st.download_button(
                    "📥 Download PDF",
                    data=pdf_bytes,
                    file_name=f"jobsuite_cover_letter.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            with dl_col2:
                st.download_button(
                    "📥 Download DOCX",
                    data=docx_bytes,
                    file_name=f"jobsuite_cover_letter.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            with dl_col3:
                st.download_button(
                    "📥 Download TXT",
                    data=txt_bytes,
                    file_name="jobsuite_cover_letter.txt",
                    mime="text/plain",
                    use_container_width=True
                )
                
        with cl_col2:
            st.markdown("**AI Co-pilot Refinement Chat**")
            
            # Show chat history
            for msg in st.session_state.cover_letter_chat:
                style_class = "chat-bubble-user" if msg["role"] == "user" else "chat-bubble-assistant"
                role_label = "**You**" if msg["role"] == "user" else "**AI Co-pilot**"
                st.markdown(f"""
                    <div class="{style_class}">
                        <div>{role_label}</div>
                        <div style="font-size:0.9rem; margin-top:4px;">{msg['content']}</div>
                    </div>
                """, unsafe_allow_html=True)
                
            # Chat input
            chat_input = st.text_input("Instruct AI on how to rewrite or refine (e.g. 'shorten the 3rd paragraph'):", key="cl_chat_input")
            if st.button("Send Instructions", key="cl_chat_send_btn"):
                if chat_input.strip():
                    st.session_state.cover_letter_chat.append({"role": "user", "content": chat_input})
                    
                    with st.spinner("AI refining letter..."):
                        refined = st.session_state.analyzer.llm_handler.refine_cover_letter(
                            st.session_state.generated_cover_letter,
                            chat_input,
                            st.session_state.cover_letter_chat
                        )
                        st.session_state.generated_cover_letter = refined
                        st.session_state.cover_letter_chat.append({
                            "role": "assistant", 
                            "content": f"I've updated the letter as requested! See the changes in the editor."
                        })
                        st.experimental_rerun()


# ==========================================
# MODULE 4: JOB BOARD CRM TRACKER
# ==========================================
elif "Job Board CRM Tracker" in nav_page:
    st.markdown('<div class="title-gradient">📊 Job Board CRM Tracker</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitle">Track companies, referral contacts, target compensation, application statuses, and interview tasks</div>', unsafe_allow_html=True)
    
    # Load applications
    applications = DBHandler.get_applications()
    
    # Columns list
    statuses = ["Wishlist", "Applied", "Interviewing", "Offer", "Rejected"]
    
    # Add application form (expander)
    with st.expander("➕ Log New Job Application", expanded=False):
        with st.form("new_app_form", clear_on_submit=True):
            app_company = st.text_input("Company Name*:")
            app_role = st.text_input("Job Title / Role*:")
            app_status = st.selectbox("Current Status:", statuses, index=0)
            app_salary = st.text_input("Expected/Offered Salary (Optional):")
            app_date = st.date_input("Date Applied / Target Date:")
            app_notes = st.text_area("Notes / Description:")
            
            submitted = st.form_submit_button("Add Application")
            if submitted:
                if not app_company or not app_role:
                    st.error("Company Name and Job Title are required.")
                else:
                    new_app = {
                        "company_name": app_company,
                        "job_title": app_role,
                        "status": app_status,
                        "salary": app_salary,
                        "date_applied": app_date.strftime("%Y-%m-%d"),
                        "notes": app_notes,
                        "contacts": [],
                        "checklist": []
                    }
                    DBHandler.add_application(new_app)
                    st.success(f"Added application for {app_role} at {app_company}!")
                    st.experimental_rerun()
                    
    # Render Kanban-like columns layout using Streamlit columns
    st.markdown("### Job search board")
    
    cols = st.columns(len(statuses))
    for col_idx, status_type in enumerate(statuses):
        with cols[col_idx]:
            # Header
            badge_class = f"status-{status_type.lower()}"
            st.markdown(f"""
                <div style="margin-bottom: 12px;">
                    <span class="kanban-status-badge {badge_class}">{status_type}</span>
                </div>
            """, unsafe_allow_html=True)
            
            # Find apps matching status
            matching_apps = [app for app in applications if app.get("status") == status_type]
            
            if not matching_apps:
                st.markdown('<div style="color:#64748b; font-size:0.85rem; font-style:italic;">Empty column</div>', unsafe_allow_html=True)
                
            for app in matching_apps:
                app_id = app["id"]
                with st.container():
                    st.markdown(f"""
                        <div class="kanban-card">
                            <strong style="color:#f1f5f9; font-size:0.95rem;">{app['job_title']}</strong><br/>
                            <span style="color:#94a3b8; font-size:0.85rem;">{app['company_name']}</span>
                            <div style="font-size:0.75rem; color:#64748b; margin-top:6px;">Applied: {app['date_applied']}</div>
                        </div>
                    """, unsafe_allow_html=True)
                    
                    # Manage panel
                    with st.expander("🛠️ Manage Details", expanded=False):
                        # Edit status dropdown
                        new_status = st.selectbox(
                            "Status:", 
                            statuses, 
                            index=statuses.index(status_type),
                            key=f"status_select_{app_id}"
                        )
                        if new_status != status_type:
                            DBHandler.update_application(app_id, {"status": new_status})
                            st.experimental_rerun()
                            
                        # Expected salary
                        salary_val = st.text_input("Salary:", value=app.get("salary", ""), key=f"salary_{app_id}")
                        if salary_val != app.get("salary"):
                            DBHandler.update_application(app_id, {"salary": salary_val})
                            
                        # Manage Notes
                        notes_val = st.text_area("Notes:", value=app.get("notes", ""), height=80, key=f"notes_{app_id}")
                        if notes_val != app.get("notes"):
                            DBHandler.update_application(app_id, {"notes": notes_val})
                            
                        # Checklist items
                        st.markdown("**Tasks / Interview prep checklist:**")
                        checklist = app.get("checklist", [])
                        
                        for task_idx, task_item in enumerate(checklist):
                            task_chk = st.checkbox(
                                task_item.get("task", ""), 
                                value=task_item.get("completed", False),
                                key=f"chk_{app_id}_{task_idx}"
                            )
                            if task_chk != task_item.get("completed"):
                                checklist[task_idx]["completed"] = task_chk
                                DBHandler.update_application(app_id, {"checklist": checklist})
                                st.experimental_rerun()
                                
                        new_task = st.text_input("Add task:", placeholder="e.g. Schedule call", key=f"new_task_{app_id}")
                        if st.button("➕ Task", key=f"btn_task_{app_id}"):
                            if new_task.strip():
                                checklist.append({"task": new_task, "completed": False})
                                DBHandler.update_application(app_id, {"checklist": checklist})
                                st.experimental_rerun()
                                
                        # Delete application button
                        st.markdown("---")
                        if st.button("🗑️ Delete Application", key=f"del_app_{app_id}"):
                            DBHandler.delete_application(app_id)
                            st.success("Deleted application.")
                            st.experimental_rerun()


# ==========================================
# MODULE 5: INTERVIEW PREP & MOCK INTERVIEW
# ==========================================
elif "Interview Prep & Mock Chat" in nav_page:
    st.markdown('<div class="title-gradient">🧠 Interview Prep & Mock Chat</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitle">Extract common, technical, and behavioral interview questions and try the interactive AI Mock Interview Simulator</div>', unsafe_allow_html=True)
    
    # Dependencies
    if not st.session_state.active_resume_text:
        st.warning("⚠️ Please upload a resume in the Optimizer or build one in the Resume Builder first to get personalized interview preparation!")
        
    jd_input = st.session_state.get('target_job_desc', "")
    jd_val = st.text_area("Target Job Requirements:", value=jd_input, height=120, placeholder="Paste job requirements here...")
    
    col_prep1, col_prep2 = st.columns(2)
    with col_prep1:
        prep_btn = st.button("✨ Generate Personalized Prep Guide", use_container_width=True)
    with col_prep2:
        mock_init_btn = st.button("🎙️ Initialize Mock Interview Simulator", use_container_width=True)
        
    # Generate prep guide logic
    if prep_btn:
        if not st.session_state.active_resume_text:
            st.error("Please provide your resume details.")
        elif not jd_val.strip():
            st.error("Please enter a target Job Description.")
        else:
            with st.spinner("AI parsing resume & job details to construct study guide..."):
                st.session_state.interview_prep_data = st.session_state.analyzer.llm_handler.generate_interview_prep(
                    st.session_state.active_resume_text,
                    jd_val
                )
                st.toast("✅ Tailored guide generated successfully!")
                
    # Display prep guide
    if st.session_state.interview_prep_data:
        prep = st.session_state.interview_prep_data
        
        st.markdown("---")
        st.markdown("### 📋 Study Guide Details")
        
        ptab1, ptab2, ptab3 = st.tabs([
            "💬 Common Questions & Answers",
            "🚀 STAR Behavioral Strategy",
            "🛠️ Technical Concepts to Master"
        ])
        
        with ptab1:
            for i, item in enumerate(prep.get("common_questions", [])):
                st.markdown(f"**Q{i+1}: {item.get('question')}**")
                st.markdown(f"*Why they ask:* {item.get('why_they_ask')}")
                st.markdown(f"*Suggested Answer Strategy:* {item.get('suggested_answer')}")
                st.markdown("---")
                
        with ptab2:
            st.markdown("Use the **STAR** framework (Situation, Task, Action, Result) to map your answers to questions.")
            for i, item in enumerate(prep.get("behavioral_star", [])):
                st.markdown(f"**Scenario {i+1}: {item.get('question')}**")
                st.markdown(f"- **Situation:** {item.get('situation')}")
                st.markdown(f"- **Task:** {item.get('task')}")
                st.markdown(f"- **Action:** {item.get('action')}")
                st.markdown(f"- **Result:** {item.get('result')}")
                st.markdown("---")
                
        with ptab3:
            for item in prep.get("technical_topics", []):
                st.markdown(f"**Topic: {item.get('topic')}**")
                st.markdown(f"*Sample Question:* {item.get('question')}")
                st.markdown("*Key Talking Points to Mention:*")
                for pt in item.get("key_talking_points", []):
                    st.markdown(f"  - {pt}")
                st.markdown("---")
                
    # Initialize mock chat simulator logic
    if mock_init_btn:
        if not jd_val.strip():
            st.error("Please enter a target Job Description for the interviewer context.")
        else:
            st.session_state.mock_last_question = "Hello! Welcome to the interview. To start off, could you introduce yourself and briefly summarize why you are interested in this position?"
            st.session_state.mock_chat = [
                {"role": "assistant", "content": st.session_state.mock_last_question}
            ]
            st.toast("🎙️ Simulator initialized!")
            st.experimental_rerun()
            
    # Mock interview session preview
    if st.session_state.mock_chat:
        st.markdown("---")
        st.markdown("### 🎙️ Mock Interview Simulator (In Progress)")
        
        # Display chat history
        for msg in st.session_state.mock_chat:
            style_class = "chat-bubble-user" if msg["role"] == "user" else "chat-bubble-assistant"
            role_label = "**You**" if msg["role"] == "user" else "**AI Interviewer**"
            st.markdown(f"""
                <div class="{style_class}">
                    <div>{role_label}</div>
                    <div style="font-size:0.9rem; margin-top:4px; white-space: pre-wrap;">{msg['content']}</div>
                </div>
            """, unsafe_allow_html=True)
            
        # Chat input
        user_answer = st.text_input("Type your interview response here:", key="mock_chat_input")
        if st.button("Submit Response", key="mock_chat_submit"):
            if user_answer.strip():
                st.session_state.mock_chat.append({"role": "user", "content": user_answer})
                
                with st.spinner("AI Interviewer evaluating response..."):
                    reply = st.session_state.analyzer.llm_handler.interview_mock_response(
                        st.session_state.mock_chat[:-1],
                        user_answer,
                        st.session_state.mock_last_question,
                        jd_val
                    )
                    
                    # Extract the next question from the AI text (which naturally asks it at the end)
                    # We store the reply in history and treat the reply as the new last question
                    st.session_state.mock_last_question = reply
                    st.session_state.mock_chat.append({"role": "assistant", "content": reply})
                    st.experimental_rerun()
                    
        if st.button("Reset Interview Simulator", key="mock_reset_btn"):
            st.session_state.mock_chat = []
            st.session_state.mock_last_question = ""
            st.experimental_rerun()
